from fastapi import FastAPI, APIRouter, HTTPException, Depends, Request, Response
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from starlette.middleware.gzip import GZipMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, EmailStr
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta
import jwt
import bcrypt
import google.generativeai as genai
import stripe as stripe_lib
import asyncio


# ============== GEMINI AI CONFIGURATION ==============
#
# Model selection notes
# ----------------------
# Gemini 2.5 Flash uses internal "thinking tokens" that count against
# max_output_tokens and that can SILENTLY consume the whole budget,
# returning a response with finish_reason=MAX_TOKENS and zero visible text.
#
# We therefore prefer `gemini-flash-latest` (which currently resolves to a
# non-thinking-by-default Flash model and accepts a much wider output
# budget) for production interpretations, and keep `gemini-2.5-flash` as
# a fallback only.
GEMINI_PRIMARY_MODEL = "gemini-flash-latest"
GEMINI_FALLBACK_MODEL = "gemini-2.5-flash"

GEMINI_DEEP_CONFIG = {
    "temperature": 0.95,
    "top_p": 0.95,
    "top_k": 40,
    "max_output_tokens": 8192,
}

GEMINI_DIRECT_CONFIG = {
    "temperature": 0.85,
    "top_p": 0.92,
    "top_k": 40,
    "max_output_tokens": 4096,
}

# Permissive safety settings: I Ching interpretations discuss life choices,
# relationships, death, spiritual symbols — all of which can trigger overly
# strict default filters. We still block CSAM and dangerous instructions.
GEMINI_SAFETY_SETTINGS = [
    {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_ONLY_HIGH"},
    {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_ONLY_HIGH"},
    {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_ONLY_HIGH"},
    {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_ONLY_HIGH"},
]


def _local_fallback_interpretation(
    primary, derived, primary_extended, derived_extended,
    hexagram_data, name_key, language, mode="direct",
) -> str:
    """
    Last-resort interpretation assembled from the traditional Wilhelm text
    we already have locally (no AI call). Better than the bare
    'not available' message — the user still gets Sentenza, Immagine,
    every active moving line, and the derived hexagram if any.
    """
    name = primary.get(name_key) or primary.get("name", "")
    giudizio = (primary_extended or {}).get("giudizio") or ""
    immagine = (primary_extended or {}).get("immagine") or ""

    moving_lines = (hexagram_data or {}).get("moving_lines") or []
    primary_n = (hexagram_data or {}).get("primary_hexagram")
    derived_n = (hexagram_data or {}).get("derived_hexagram")

    if language == "en":
        parts = [
            f"### {name} (#{primary_n})",
            "",
            "**The Judgment**",
            giudizio.strip() or "—",
            "",
            "**The Image**",
            immagine.strip() or "—",
        ]
    else:
        parts = [
            f"### {name} (#{primary_n})",
            "",
            "**La Sentenza (Giudizio)**",
            giudizio.strip() or "—",
            "",
            "**L'Immagine**",
            immagine.strip() or "—",
        ]

    if moving_lines:
        parts.append("")
        parts.append("**Linee mutevoli**" if language != "en" else "**Moving lines**")
        for pos in moving_lines:
            try:
                ld = get_moving_line_extended(primary_n, pos, language) or {}
                testo = (ld.get("testo") or ld.get("text") or "").strip()
                spieg = (ld.get("commento") or ld.get("commentary") or ld.get("interpretation") or "").strip()
                if testo:
                    parts.append(f"- **{('Linea' if language != 'en' else 'Line')} {pos}.** {testo}")
                if spieg:
                    parts.append(f"  {spieg}")
            except Exception:
                continue

    if derived_n and derived:
        d_name = derived.get(name_key) or derived.get("name", "")
        d_giud = (derived_extended or {}).get("giudizio") or ""
        parts.append("")
        if language == "en":
            parts.append(f"**Where it is heading → {d_name} (#{derived_n})**")
        else:
            parts.append(f"**Verso cui si dirige → {d_name} (#{derived_n})**")
        if d_giud:
            parts.append(d_giud.strip())

    parts.append("")
    if language == "en":
        parts.append(
            "_(Service note: AI elaboration is momentarily unavailable. "
            "The text above is the traditional reading from Richard Wilhelm's translation.)_"
        )
    else:
        parts.append(
            "_(Nota di servizio: l'elaborazione AI non è momentaneamente disponibile. "
            "Il testo qui sopra è la lettura tradizionale tratta dalla traduzione di Richard Wilhelm.)_"
        )
    return "\n".join(parts)


def _extract_text_from_gemini_response(response) -> str:
    """
    Robust text extraction from a Gemini response.

    `response.text` is a *property* in google-generativeai that raises
    ValueError when finish_reason is not STOP (e.g. MAX_TOKENS, SAFETY,
    or when there is a `thoughtSignature` part but no text). We must
    therefore NEVER touch `.text` blindly — even `getattr(..., "text", None)`
    actually invokes the property and triggers the exception.

    Strategy:
      1) Try `.text` inside a try/except. If it works, perfect.
      2) Otherwise iterate `candidates -> content.parts` and concatenate
         every `.text` we find on the parts (skipping `thought` parts).
    """
    if not response:
        return ""
    # 1. Fast path
    try:
        t = response.text
        if t:
            return t
    except Exception:
        pass
    # 2. Manual extraction from candidates/parts
    pieces = []
    try:
        candidates = getattr(response, "candidates", None) or []
        for cand in candidates:
            content = getattr(cand, "content", None)
            if not content:
                continue
            for part in getattr(content, "parts", []) or []:
                # Skip Gemini 2.5 'thought' parts that carry no user-facing text
                if getattr(part, "thought", False):
                    continue
                ptext = getattr(part, "text", None)
                if ptext:
                    pieces.append(ptext)
    except Exception:
        pass
    return "".join(pieces)


async def _gemini_generate_with_retry(model, prompt, max_retries: int = 2):
    """
    Call ONE Gemini model with exponential backoff on rate-limit / 5xx.
    Returns the extracted text (possibly empty).
    """
    last_error = None
    for attempt in range(max_retries):
        try:
            response = await model.generate_content_async(prompt)
            text = _extract_text_from_gemini_response(response)
            if text:
                return text
            try:
                fr = response.candidates[0].finish_reason if response.candidates else "?"
            except Exception:
                fr = "?"
            logger.warning(f"Gemini empty text (finish_reason={fr}, attempt {attempt+1}/{max_retries})")
            return ""
        except Exception as e:
            last_error = e
            err_str = str(e).lower()
            if any(s in err_str for s in ("429", "rate", "quota", "timeout", "503", "504", "500")):
                wait = 2 ** attempt
                logger.warning(f"Gemini transient error (attempt {attempt+1}/{max_retries}), retrying in {wait}s: {e}")
                await asyncio.sleep(wait)
                continue
            logger.error(f"Gemini non-transient error: {e}", exc_info=True)
            raise
    raise last_error if last_error else RuntimeError("Gemini retry exhausted")


async def _gemini_generate(
    *,
    system_instruction: str,
    prompt: str,
    generation_config: dict,
    primary_model: str = None,
    fallback_model: str = None,
):
    """
    Try the primary Gemini model first; if it returns empty text (or fails
    with a transient error), automatically retry on a different model.
    Returns text (possibly empty). All Gemini-side errors are caught here
    so the caller can rely on a string return.
    """
    primary_model = primary_model or GEMINI_PRIMARY_MODEL
    fallback_model = fallback_model or GEMINI_FALLBACK_MODEL

    for model_name in (primary_model, fallback_model):
        try:
            model = genai.GenerativeModel(
                model_name=model_name,
                system_instruction=system_instruction,
                generation_config=generation_config,
                safety_settings=GEMINI_SAFETY_SETTINGS,
            )
            text = await _gemini_generate_with_retry(model, prompt)
            if text and len(text) > 100:
                if model_name != primary_model:
                    logger.info(f"Gemini used FALLBACK model {model_name}")
                return text
            logger.warning(f"Gemini model {model_name} returned no useful text — trying next")
        except Exception as e:
            logger.error(f"Gemini model {model_name} raised: {e}")
            continue
    return ""
from iching_data import get_hexagram_traditional_data, get_trigram_info, get_moving_lines_text, get_all_lines_text, TRIGRAMS
from iching_extended import ICHING_EXTENDED, get_extended_hexagram_data, get_moving_line_extended
from subscription_manager import (
    get_user_plan, get_plan_limits, check_consultation_limit, can_use_consultation_type,
    get_daily_hexagram_number, get_lunar_phase, get_user_level, check_and_award_badges,
    is_admin_email,
    grant_trial_pack, consume_trial_credit_if_applicable,
    PLAN_LIMITS, SUBSCRIPTION_PRICES, USER_LEVELS, BADGES, GUIDED_PATHS
)
import email_service as mailer
from personalized_advice import (
    generate_personalized_advice, get_chinese_day_energy, get_chinese_year_animal,
    get_user_notification_preferences, update_user_notification_preferences
)
from astrology_profile import (
    get_full_astrological_profile, validate_profile_data, USER_PROFILE_FIELDS,
    calculate_chinese_zodiac, calculate_western_zodiac
)
from natal_chart import calculate_natal_chart, geocode_location, KERYKEION_AVAILABLE
from wilhelm_source import build_authoritative_context, is_loaded as wilhelm_loaded
from fitness_coaching import (
    ONBOARDING_QUESTIONS as FITNESS_ONBOARDING_QUESTIONS,
    validate_onboarding as fitness_validate_onboarding,
    score_to_focus_areas as fitness_score,
    generate_weekly_program as fitness_generate_program,
    compute_xp_and_badges as fitness_compute_xp,
)

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Config
# JWT_SECRET deve essere settato come env var in produzione. Se l'env
# var manca: usiamo un default *solo* in sviluppo (e logghiamo warning
# forte). In produzione (Render espone RENDER=true) rifiutiamo di
# avviarci — meglio un crash chiaro che token forgiabili.
_JWT_DEFAULT = 'iching-secret-DEV-ONLY-not-for-production'
JWT_SECRET = os.environ.get('JWT_SECRET') or _JWT_DEFAULT
if JWT_SECRET == _JWT_DEFAULT:
    if os.environ.get('RENDER') or os.environ.get('ENV') == 'production':
        raise RuntimeError(
            "JWT_SECRET non configurato in produzione! "
            "Imposta la variabile su Render → Environment prima di riavviare. "
            "Senza JWT_SECRET unico e segreto, qualsiasi attaccante puo' "
            "forgiare token validi."
        )
    import logging as _log
    _log.warning(
        "⚠️  JWT_SECRET non configurato — sto usando un default DI SVILUPPO. "
        "Non usare in produzione."
    )
GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY')
STRIPE_API_KEY = os.environ.get('STRIPE_API_KEY')
STRIPE_WEBHOOK_SECRET = os.environ.get('STRIPE_WEBHOOK_SECRET')

# Configure Gemini
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

# Create the main app
app = FastAPI(title="I Ching del Benessere API")
api_router = APIRouter(prefix="/api")

# Logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ============== I CHING DATA ==============
HEXAGRAMS = {
    1: {"name": "乾 Qián", "name_it": "Il Creativo", "name_en": "The Creative", "trigram_top": "☰", "trigram_bottom": "☰"},
    2: {"name": "坤 Kūn", "name_it": "Il Ricettivo", "name_en": "The Receptive", "trigram_top": "☷", "trigram_bottom": "☷"},
    3: {"name": "屯 Zhūn", "name_it": "La Difficoltà Iniziale", "name_en": "Difficulty at the Beginning", "trigram_top": "☵", "trigram_bottom": "☳"},
    4: {"name": "蒙 Méng", "name_it": "La Stoltezza Giovanile", "name_en": "Youthful Folly", "trigram_top": "☶", "trigram_bottom": "☵"},
    5: {"name": "需 Xū", "name_it": "L'Attesa", "name_en": "Waiting", "trigram_top": "☵", "trigram_bottom": "☰"},
    6: {"name": "訟 Sòng", "name_it": "Il Conflitto", "name_en": "Conflict", "trigram_top": "☰", "trigram_bottom": "☵"},
    7: {"name": "師 Shī", "name_it": "L'Esercito", "name_en": "The Army", "trigram_top": "☷", "trigram_bottom": "☵"},
    8: {"name": "比 Bǐ", "name_it": "L'Unione", "name_en": "Holding Together", "trigram_top": "☵", "trigram_bottom": "☷"},
    9: {"name": "小畜 Xiǎo Chù", "name_it": "La Forza Domatrice Piccola", "name_en": "Small Taming", "trigram_top": "☴", "trigram_bottom": "☰"},
    10: {"name": "履 Lǚ", "name_it": "Il Procedere", "name_en": "Treading", "trigram_top": "☰", "trigram_bottom": "☱"},
    11: {"name": "泰 Tài", "name_it": "La Pace", "name_en": "Peace", "trigram_top": "☷", "trigram_bottom": "☰"},
    12: {"name": "否 Pǐ", "name_it": "Il Ristagno", "name_en": "Standstill", "trigram_top": "☰", "trigram_bottom": "☷"},
    13: {"name": "同人 Tóng Rén", "name_it": "La Comunità", "name_en": "Fellowship", "trigram_top": "☰", "trigram_bottom": "☲"},
    14: {"name": "大有 Dà Yǒu", "name_it": "Il Possesso Grande", "name_en": "Great Possession", "trigram_top": "☲", "trigram_bottom": "☰"},
    15: {"name": "謙 Qiān", "name_it": "La Modestia", "name_en": "Modesty", "trigram_top": "☷", "trigram_bottom": "☶"},
    16: {"name": "豫 Yù", "name_it": "L'Entusiasmo", "name_en": "Enthusiasm", "trigram_top": "☳", "trigram_bottom": "☷"},
    17: {"name": "隨 Suí", "name_it": "Il Seguire", "name_en": "Following", "trigram_top": "☱", "trigram_bottom": "☳"},
    18: {"name": "蠱 Gǔ", "name_it": "L'Emendamento", "name_en": "Work on the Decayed", "trigram_top": "☶", "trigram_bottom": "☴"},
    19: {"name": "臨 Lín", "name_it": "L'Avvicinamento", "name_en": "Approach", "trigram_top": "☷", "trigram_bottom": "☱"},
    20: {"name": "觀 Guān", "name_it": "La Contemplazione", "name_en": "Contemplation", "trigram_top": "☴", "trigram_bottom": "☷"},
    21: {"name": "噬嗑 Shì Kè", "name_it": "Il Morso", "name_en": "Biting Through", "trigram_top": "☲", "trigram_bottom": "☳"},
    22: {"name": "賁 Bì", "name_it": "L'Avvenenza", "name_en": "Grace", "trigram_top": "☶", "trigram_bottom": "☲"},
    23: {"name": "剝 Bō", "name_it": "Lo Sgretolamento", "name_en": "Splitting Apart", "trigram_top": "☶", "trigram_bottom": "☷"},
    24: {"name": "復 Fù", "name_it": "Il Ritorno", "name_en": "Return", "trigram_top": "☷", "trigram_bottom": "☳"},
    25: {"name": "無妄 Wú Wàng", "name_it": "L'Innocenza", "name_en": "Innocence", "trigram_top": "☰", "trigram_bottom": "☳"},
    26: {"name": "大畜 Dà Chù", "name_it": "La Forza Domatrice Grande", "name_en": "Great Taming", "trigram_top": "☶", "trigram_bottom": "☰"},
    27: {"name": "頤 Yí", "name_it": "Gli Angoli della Bocca", "name_en": "Nourishment", "trigram_top": "☶", "trigram_bottom": "☳"},
    28: {"name": "大過 Dà Guò", "name_it": "La Preponderanza del Grande", "name_en": "Great Excess", "trigram_top": "☱", "trigram_bottom": "☴"},
    29: {"name": "坎 Kǎn", "name_it": "L'Abissale", "name_en": "The Abysmal", "trigram_top": "☵", "trigram_bottom": "☵"},
    30: {"name": "離 Lí", "name_it": "L'Aderente", "name_en": "The Clinging", "trigram_top": "☲", "trigram_bottom": "☲"},
    31: {"name": "咸 Xián", "name_it": "L'Influsso", "name_en": "Influence", "trigram_top": "☱", "trigram_bottom": "☶"},
    32: {"name": "恆 Héng", "name_it": "La Durata", "name_en": "Duration", "trigram_top": "☳", "trigram_bottom": "☴"},
    33: {"name": "遯 Dùn", "name_it": "La Ritirata", "name_en": "Retreat", "trigram_top": "☰", "trigram_bottom": "☶"},
    34: {"name": "大壯 Dà Zhuàng", "name_it": "La Potenza del Grande", "name_en": "Great Power", "trigram_top": "☳", "trigram_bottom": "☰"},
    35: {"name": "晉 Jìn", "name_it": "Il Progresso", "name_en": "Progress", "trigram_top": "☲", "trigram_bottom": "☷"},
    36: {"name": "明夷 Míng Yí", "name_it": "L'Ottenebramento della Luce", "name_en": "Darkening of the Light", "trigram_top": "☷", "trigram_bottom": "☲"},
    37: {"name": "家人 Jiā Rén", "name_it": "La Famiglia", "name_en": "The Family", "trigram_top": "☴", "trigram_bottom": "☲"},
    38: {"name": "睽 Kuí", "name_it": "L'Opposizione", "name_en": "Opposition", "trigram_top": "☲", "trigram_bottom": "☱"},
    39: {"name": "蹇 Jiǎn", "name_it": "L'Impedimento", "name_en": "Obstruction", "trigram_top": "☵", "trigram_bottom": "☶"},
    40: {"name": "解 Xiè", "name_it": "La Liberazione", "name_en": "Deliverance", "trigram_top": "☳", "trigram_bottom": "☵"},
    41: {"name": "損 Sǔn", "name_it": "La Diminuzione", "name_en": "Decrease", "trigram_top": "☶", "trigram_bottom": "☱"},
    42: {"name": "益 Yì", "name_it": "L'Accrescimento", "name_en": "Increase", "trigram_top": "☴", "trigram_bottom": "☳"},
    43: {"name": "夬 Guài", "name_it": "L'Irrompere", "name_en": "Breakthrough", "trigram_top": "☱", "trigram_bottom": "☰"},
    44: {"name": "姤 Gòu", "name_it": "Il Farsi Incontro", "name_en": "Coming to Meet", "trigram_top": "☰", "trigram_bottom": "☴"},
    45: {"name": "萃 Cuì", "name_it": "La Raccolta", "name_en": "Gathering Together", "trigram_top": "☱", "trigram_bottom": "☷"},
    46: {"name": "升 Shēng", "name_it": "L'Ascesa", "name_en": "Pushing Upward", "trigram_top": "☷", "trigram_bottom": "☴"},
    47: {"name": "困 Kùn", "name_it": "L'Angustia", "name_en": "Oppression", "trigram_top": "☱", "trigram_bottom": "☵"},
    48: {"name": "井 Jǐng", "name_it": "Il Pozzo", "name_en": "The Well", "trigram_top": "☵", "trigram_bottom": "☴"},
    49: {"name": "革 Gé", "name_it": "La Rivoluzione", "name_en": "Revolution", "trigram_top": "☱", "trigram_bottom": "☲"},
    50: {"name": "鼎 Dǐng", "name_it": "Il Crogiolo", "name_en": "The Cauldron", "trigram_top": "☲", "trigram_bottom": "☴"},
    51: {"name": "震 Zhèn", "name_it": "L'Eccitante", "name_en": "The Arousing", "trigram_top": "☳", "trigram_bottom": "☳"},
    52: {"name": "艮 Gèn", "name_it": "L'Arresto", "name_en": "Keeping Still", "trigram_top": "☶", "trigram_bottom": "☶"},
    53: {"name": "漸 Jiàn", "name_it": "Lo Sviluppo Graduale", "name_en": "Development", "trigram_top": "☴", "trigram_bottom": "☶"},
    54: {"name": "歸妹 Guī Mèi", "name_it": "La Ragazza che si Marita", "name_en": "The Marrying Maiden", "trigram_top": "☳", "trigram_bottom": "☱"},
    55: {"name": "豐 Fēng", "name_it": "L'Abbondanza", "name_en": "Abundance", "trigram_top": "☳", "trigram_bottom": "☲"},
    56: {"name": "旅 Lǚ", "name_it": "Il Viandante", "name_en": "The Wanderer", "trigram_top": "☲", "trigram_bottom": "☶"},
    57: {"name": "巽 Xùn", "name_it": "Il Mite", "name_en": "The Gentle", "trigram_top": "☴", "trigram_bottom": "☴"},
    58: {"name": "兌 Duì", "name_it": "Il Sereno", "name_en": "The Joyous", "trigram_top": "☱", "trigram_bottom": "☱"},
    59: {"name": "渙 Huàn", "name_it": "La Dissoluzione", "name_en": "Dispersion", "trigram_top": "☴", "trigram_bottom": "☵"},
    60: {"name": "節 Jié", "name_it": "La Limitazione", "name_en": "Limitation", "trigram_top": "☵", "trigram_bottom": "☱"},
    61: {"name": "中孚 Zhōng Fú", "name_it": "La Verità Interiore", "name_en": "Inner Truth", "trigram_top": "☴", "trigram_bottom": "☱"},
    62: {"name": "小過 Xiǎo Guò", "name_it": "La Preponderanza del Piccolo", "name_en": "Small Excess", "trigram_top": "☳", "trigram_bottom": "☶"},
    63: {"name": "既濟 Jì Jì", "name_it": "Dopo il Compimento", "name_en": "After Completion", "trigram_top": "☵", "trigram_bottom": "☲"},
    64: {"name": "未濟 Wèi Jì", "name_it": "Prima del Compimento", "name_en": "Before Completion", "trigram_top": "☲", "trigram_bottom": "☵"},
}

# Binary to Hexagram mapping (bottom to top, 0=yin, 1=yang)
BINARY_TO_HEX = {
    "111111": 1, "000000": 2, "010001": 3, "100010": 4, "010111": 5, "111010": 6,
    "000010": 7, "010000": 8, "110111": 9, "111011": 10, "000111": 11, "111000": 12,
    "111101": 13, "101111": 14, "000100": 15, "001000": 16, "011001": 17, "100110": 18,
    "000011": 19, "110000": 20, "101001": 21, "100101": 22, "100000": 23, "000001": 24,
    "111001": 25, "100111": 26, "100001": 27, "011110": 28, "010010": 29, "101101": 30,
    "011100": 31, "001110": 32, "111100": 33, "001111": 34, "101000": 35, "000101": 36,
    "110101": 37, "101011": 38, "010100": 39, "001010": 40, "100011": 41, "110001": 42,
    "011111": 43, "111110": 44, "011000": 45, "000110": 46, "011010": 47, "010110": 48,
    "011101": 49, "101110": 50, "001001": 51, "100100": 52, "110100": 53, "001011": 54,
    "001101": 55, "101100": 56, "110110": 57, "011011": 58, "110010": 59, "010011": 60,
    "110011": 61, "001100": 62, "010101": 63, "101010": 64,
}

# ============== MODELS ==============
class UserCreate(BaseModel):
    email: EmailStr
    # Server-side validation: minimum 6 chars (matches frontend hint).
    # Keeps registration safe even if a client bypasses UI validation.
    password: str = Field(..., min_length=6, max_length=128)
    name: str = Field(..., min_length=1, max_length=80)
    phone: str = Field(default="", max_length=30)
    language: str = "it"
    # GDPR consents (required by art. 7 GDPR — proof of consent)
    privacy_accepted: bool = False
    marketing_consent: bool = False

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class UserResponse(BaseModel):
    id: str
    email: str
    name: str
    phone: str = ""
    language: str
    subscription_active: bool = False
    subscription_end: Optional[str] = None
    is_admin: bool = False  # True for owner/whitelisted emails
    plan: str = "free"      # 'premium' for admins & paying users

class PasswordResetRequest(BaseModel):
    email: EmailStr
    phone: str = ""

class PasswordResetVerify(BaseModel):
    email: EmailStr
    code: str
    new_password: str

class CoinToss(BaseModel):
    line1: int = Field(..., ge=6, le=9)
    line2: int = Field(..., ge=6, le=9)
    line3: int = Field(..., ge=6, le=9)
    line4: int = Field(..., ge=6, le=9)
    line5: int = Field(..., ge=6, le=9)
    line6: int = Field(..., ge=6, le=9)

class ConsultationCreate(BaseModel):
    question: str
    coin_tosses: CoinToss
    consultation_type: str = "deep"  # "direct" or "deep"
    parent_consultation_id: Optional[str] = None  # For continuing a conversation
    topic: Optional[str] = None  # 'amore', 'lavoro', 'fortuna', 'soldi', 'spirituale', 'personale', or custom text

class TrigramInfo(BaseModel):
    symbol: str
    name: str
    name_local: str
    element: str
    quality: str
    color: str

class MovingLineText(BaseModel):
    position: int
    text: str
    meaning: str
    is_active: bool = True  # True if this line is a moving line

class TraditionalData(BaseModel):
    sentence: str
    image: str
    commentary: str = ""
    trigram_above: TrigramInfo
    trigram_below: TrigramInfo
    moving_lines_text: List[MovingLineText]

class ConsultationResponse(BaseModel):
    id: str
    question: str
    hexagram_number: int
    hexagram_name: str
    hexagram_chinese: str
    hexagram_symbol: str
    derived_hexagram_number: Optional[int] = None
    derived_hexagram_name: Optional[str] = None
    derived_hexagram_chinese: Optional[str] = None
    moving_lines: List[int]
    traditional_data: Optional[TraditionalData] = None
    derived_traditional_data: Optional[TraditionalData] = None
    interpretation: str
    created_at: str
    consultation_type: str = "deep"  # "direct" or "deep"
    # Fields for conversation/continuation
    parent_consultation_id: Optional[str] = None
    conversation_depth: int = 0  # How many consultations deep in the conversation
    # Fields for linked consultations (synthesis)
    is_synthesis: bool = False
    linked_consultation_ids: List[str] = []
    synthesis_type: Optional[str] = None  # "confirmation", "deepening", "clarification"

class SynthesisRequest(BaseModel):
    consultation_ids: List[str]
    synthesis_type: str = "deepening"  # confirmation, deepening, clarification

class CheckoutRequest(BaseModel):
    origin_url: str
    # Accepted values:
    #   base_monthly      -> €9,99 / 30gg
    #   base_yearly       -> €107,89 / 365gg  (-10%)
    #   fitness_monthly   -> €19,99 / 30gg
    #   fitness_yearly    -> €191,90 / 365gg  (-20%)
    # Legacy aliases still accepted: 'monthly' -> base_monthly,
    # 'yearly' -> base_yearly.
    plan_type: str = "base_monthly"

class NoteCreate(BaseModel):
    consultation_id: str
    content: str
    mood: Optional[str] = None  # 'positive', 'neutral', 'negative', 'reflective'
    tags: Optional[List[str]] = []

class NoteUpdate(BaseModel):
    content: Optional[str] = None
    mood: Optional[str] = None
    tags: Optional[List[str]] = None

# ============== AUTH HELPERS ==============
import random
import string
import secrets
import hmac

# Rate limiting per endpoint sensibili (brute-force, abuso AI, DoS).
# Vedi requirements.txt: slowapi==0.1.9
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded


def _client_ip(request) -> str:
    """
    Estrae l'IP client reale, non l'IP del proxy Render/Vercel/Cloudflare.
    Senza questo, `get_remote_address` restituisce sempre l'IP del reverse
    proxy Render — cosi' il bucket rate-limit e' UNICO PER TUTTO IL MONDO:
    un utente che fa login triggera il limite anche per gli altri, e un
    attaccante non viene mai davvero bloccato.

    Priorita': X-Forwarded-For (primo IP della catena) -> X-Real-IP ->
    request.client.host. Il primo IP di XFF e' il client originale, gli
    altri sono i proxy intermedi.
    """
    xff = request.headers.get("x-forwarded-for", "")
    if xff:
        # "client, proxy1, proxy2" -> "client"
        first = xff.split(",")[0].strip()
        if first:
            return first
    xri = request.headers.get("x-real-ip", "")
    if xri:
        return xri.strip()
    return get_remote_address(request)


limiter = Limiter(key_func=_client_ip)

def generate_reset_code():
    """
    Genera un codice di reset a 8 cifre con un CSPRNG (secrets.choice),
    non con random.choices(). Motivo:
      - random.choices usa il Mersenne Twister, predicibile osservando
        abbastanza output → un attaccante che chiede molti reset puo'
        prevedere i futuri codici.
      - 6 cifre = 10^6 combinazioni → bruteforciabile in pochi minuti
        senza rate-limiting. 8 cifre = 10^8 → 100x piu' costoso.
    L'uso di hmac.compare_digest in verify_reset_code() chiude il
    timing channel.
    """
    return ''.join(secrets.choice(string.digits) for _ in range(8))

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode(), hashed.encode())

def create_token(user_id: str, email: str) -> str:
    payload = {
        "user_id": user_id,
        "email": email,
        "exp": datetime.now(timezone.utc) + timedelta(days=30)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm="HS256")

async def get_current_user(request: Request) -> dict:
    auth_header = request.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Non autenticato")
    token = auth_header.split(" ")[1]
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
        user = await db.users.find_one({"id": payload["user_id"]}, {"_id": 0})
        if not user:
            raise HTTPException(status_code=401, detail="Utente non trovato")
        return user
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token scaduto")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Token non valido")

# ============== I CHING LOGIC ==============
def calculate_hexagram(coin_tosses: CoinToss) -> dict:
    # Lines in order: line1 (bottom) to line6 (top)
    lines = [coin_tosses.line1, coin_tosses.line2, coin_tosses.line3, 
             coin_tosses.line4, coin_tosses.line5, coin_tosses.line6]
    
    # Convert to binary - IMPORTANT: The binary string must be built from LINE 6 (top) to LINE 1 (bottom)
    # because the BINARY_TO_HEX mapping expects: first char = line 6 (top), last char = line 1 (bottom)
    primary_binary = ""
    derived_binary = ""
    moving_lines = []
    
    # Process lines in REVERSE order (from line 6 down to line 1) to build correct binary string
    for i in range(5, -1, -1):  # 5, 4, 3, 2, 1, 0 -> line 6, 5, 4, 3, 2, 1
        line = lines[i]
        line_position = i + 1  # 1-based position
        
        if line == 6:  # Old Yin (mutevole) - transforms to Yang
            primary_binary += "0"
            derived_binary += "1"
            moving_lines.append(line_position)
        elif line == 7:  # Young Yang
            primary_binary += "1"
            derived_binary += "1"
        elif line == 8:  # Young Yin
            primary_binary += "0"
            derived_binary += "0"
        elif line == 9:  # Old Yang (mutevole) - transforms to Yin
            primary_binary += "1"
            derived_binary += "0"
            moving_lines.append(line_position)
    
    # Sort moving lines in ascending order (1 to 6)
    moving_lines.sort()
    
    primary_hex = BINARY_TO_HEX.get(primary_binary, 1)
    derived_hex = BINARY_TO_HEX.get(derived_binary, None) if moving_lines else None
    
    return {
        "primary_hexagram": primary_hex,
        "derived_hexagram": derived_hex,
        "moving_lines": moving_lines,
        "lines": lines
    }

def get_hexagram_symbol(lines: List[int]) -> str:
    """Generate hexagram symbol from lines"""
    symbols = []
    for line in lines:
        if line in [7, 9]:  # Yang
            symbols.append("━━━━━" if line == 7 else "━━○━━")
        else:  # Yin
            symbols.append("━━ ━━" if line == 8 else "━━×━━")
    return "\n".join(reversed(symbols))

async def get_conversation_history(parent_id: str, user_id: str, max_depth: int = 5) -> list:
    """Retrieve the conversation history by following parent_consultation_id chain"""
    history = []
    current_id = parent_id
    depth = 0
    
    while current_id and depth < max_depth:
        consultation = await db.consultations.find_one(
            {"id": current_id, "user_id": user_id},
            {"_id": 0, "question": 1, "hexagram_number": 1, "hexagram_name": 1, 
             "interpretation": 1, "parent_consultation_id": 1, "moving_lines": 1,
             "derived_hexagram_number": 1, "derived_hexagram_name": 1}
        )
        
        if not consultation:
            break
            
        history.insert(0, consultation)  # Insert at beginning to maintain chronological order
        current_id = consultation.get("parent_consultation_id")
        depth += 1
    
    return history

async def generate_interpretation(hexagram_data: dict, question: str, language: str, 
                                   consultation_type: str = "deep", 
                                   conversation_history: list = None,
                                   topic: str = None) -> str:
    """Generate AI interpretation using Gemini - either direct or deep style"""
    primary = HEXAGRAMS.get(hexagram_data["primary_hexagram"], {})
    derived = HEXAGRAMS.get(hexagram_data["derived_hexagram"], {}) if hexagram_data["derived_hexagram"] else None
    
    # Get extended data from the Book of Changes
    primary_extended = get_extended_hexagram_data(hexagram_data["primary_hexagram"], language)
    derived_extended = get_extended_hexagram_data(hexagram_data["derived_hexagram"], language) if hexagram_data["derived_hexagram"] else None
    
    name_key = "name_it" if language == "it" else "name_en"
    
    # DIRECT STYLE - Simple, impactful, to the point
    if consultation_type == "direct":
        return await generate_direct_interpretation(
            hexagram_data, question, language, primary, derived, 
            primary_extended, derived_extended, name_key,
            conversation_history=conversation_history,
            topic=topic
        )
    
    # PATH STYLE - For guided path consultations (uses deep style with path context)
    is_path_consultation = consultation_type == "path"
    
    # DEEP STYLE - Full traditional interpretation with Book of Changes quotes
    
    # Topic context for more focused interpretations
    topic_context_it = ""
    topic_context_en = ""
    if topic:
        topic_map_it = {
            'amore': 'AMORE E RELAZIONI - Focalizza l\'interpretazione su aspetti sentimentali, relazioni di coppia, affetti familiari, amicizie profonde',
            'lavoro': 'LAVORO E CARRIERA - Focalizza l\'interpretazione su aspetti professionali, carriera, progetti lavorativi, rapporti con colleghi e superiori',
            'fortuna': 'FORTUNA E OPPORTUNITÀ - Focalizza l\'interpretazione su opportunità future, eventi favorevoli, destino, tempismo delle azioni',
            'soldi': 'FINANZE E DENARO - Focalizza l\'interpretazione su aspetti economici, investimenti, prosperità materiale, gestione delle risorse',
            'spirituale': 'CRESCITA SPIRITUALE - Focalizza l\'interpretazione su evoluzione interiore, ricerca spirituale, meditazione, connessione con il Tao',
            'personale': 'CRESCITA PERSONALE - Focalizza l\'interpretazione su sviluppo personale, raggiungimento obiettivi, superamento limiti, miglioramento di sé'
        }
        topic_map_en = {
            'amore': 'LOVE AND RELATIONSHIPS - Focus the interpretation on romantic aspects, couple relationships, family bonds, deep friendships',
            'lavoro': 'WORK AND CAREER - Focus the interpretation on professional aspects, career, work projects, relationships with colleagues and superiors',
            'fortuna': 'FORTUNE AND OPPORTUNITIES - Focus the interpretation on future opportunities, favorable events, destiny, timing of actions',
            'soldi': 'FINANCES AND MONEY - Focus the interpretation on economic aspects, investments, material prosperity, resource management',
            'spirituale': 'SPIRITUAL GROWTH - Focus the interpretation on inner evolution, spiritual search, meditation, connection with the Tao',
            'personale': 'PERSONAL GROWTH - Focus the interpretation on personal development, achieving goals, overcoming limits, self-improvement'
        }
        topic_context_it = topic_map_it.get(topic, f'ARGOMENTO SPECIFICO: {topic} - Focalizza l\'interpretazione su questo tema specifico indicato dal consultante')
        topic_context_en = topic_map_en.get(topic, f'SPECIFIC TOPIC: {topic} - Focus the interpretation on this specific topic indicated by the querent')
    
    if language == "it":
        topic_instruction = f"\n\n**ARGOMENTO DELLA DOMANDA:**\n{topic_context_it}\nDevi interpretare OGNI aspetto dell'esagramma in relazione a questo argomento specifico. Sii CONCRETO e PRATICO nei consigli relativi a questo tema." if topic_context_it else ""
        
        system_prompt = f"""Sei un venerabile Maestro dell'I Ching, custode della saggezza millenaria del Libro dei Mutamenti.
La tua voce è quella di un antico saggio taoista che parla con profondità, poesia e compassione.
{topic_instruction}

STILE DI SCRITTURA:
- Scrivi in modo contemplativo, evocativo e profondamente spirituale
- Usa metafore dalla natura: acqua che scorre, montagne, vento, stagioni, draghi, tigri
- Parla SEMPRE in seconda persona al consultante ("tu", "il tuo cammino", "la tua domanda")
- Mai elenchi puntati o strutture meccaniche - solo prosa fluida e narrativa
- Evoca immagini visive e sensoriali
- La risposta deve suonare come se fosse pronunciata da un maestro millenario

STRUTTURA DELL'INTERPRETAZIONE (senza titoli espliciti, tutto fluido):
1. Apertura poetica che connette il consultante al flusso del Tao
2. Spiegazione profonda dell'esagramma principale con i suoi trigrammi
3. Il Giudizio e l'Immagine tradizionali spiegati in relazione alla domanda
4. Se ci sono LINEE MUTEVOLI: spiegazione DETTAGLIATA di OGNI linea che muta, con il testo tradizionale e il suo significato profondo applicato alla situazione
5. Se c'è esagramma DERIVATO: spiegazione della trasformazione e del suo significato
6. Conclusione con consiglio pratico e saggezza applicabile

IMPORTANTE:
- Ogni interpretazione deve essere UNICA e SPECIFICA per la domanda posta
- Cita i testi tradizionali quando appropriato (il Giudizio, l'Immagine)
- Le linee mutevoli sono CRUCIALI - dedica almeno un paragrafo a ciascuna
- L'esagramma derivato indica DOVE si sta andando - spiegalo chiaramente
- Lunghezza: 600-900 parole per un'interpretazione completa e soddisfacente"""
    else:
        topic_instruction = f"\n\n**QUESTION TOPIC:**\n{topic_context_en}\nYou must interpret EVERY aspect of the hexagram in relation to this specific topic. Be CONCRETE and PRACTICAL in your advice related to this theme." if topic_context_en else ""
        
        system_prompt = f"""You are a venerable Master of the I Ching, guardian of the ancient wisdom of the Book of Changes.
Your voice is that of an ancient Taoist sage who speaks with depth, poetry, and compassion.
{topic_instruction}

WRITING STYLE:
- Write contemplatively, evocatively, and deeply spiritually
- Use metaphors from nature: flowing water, mountains, wind, seasons, dragons, tigers
- ALWAYS speak in second person to the querent ("you", "your path", "your question")
- Never bullet points or mechanical structures - only fluid, narrative prose
- Evoke visual and sensory images
- The response should sound as if spoken by an ancient master

INTERPRETATION STRUCTURE (no explicit titles, all flowing):
1. Poetic opening connecting the querent to the flow of Tao
2. Deep explanation of the primary hexagram with its trigrams
3. The traditional Judgment and Image explained in relation to the question
4. If there are MOVING LINES: DETAILED explanation of EACH changing line, with traditional text and deep meaning applied to the situation
5. If there is a DERIVED hexagram: explanation of the transformation and its meaning
6. Conclusion with practical advice and applicable wisdom

IMPORTANT:
- Each interpretation must be UNIQUE and SPECIFIC to the question asked
- Quote traditional texts when appropriate (the Judgment, the Image)
- Moving lines are CRUCIAL - dedicate at least one paragraph to each
- The derived hexagram indicates WHERE things are going - explain clearly
- Length: 600-900 words for a complete and satisfying interpretation"""

    # Build the detailed context for the AI
    primary_name = primary.get(name_key, primary.get("name", ""))
    primary_chinese = primary.get("name", "")
    
    # Extended data
    giudizio = primary_extended.get("giudizio", "")
    immagine = primary_extended.get("immagine", "")
    commento = primary_extended.get("commento", "")
    trigramma_sup = primary_extended.get("trigramma_superiore", primary.get("trigram_top", ""))
    trigramma_inf = primary_extended.get("trigramma_inferiore", primary.get("trigram_bottom", ""))
    
    # Moving lines details
    moving_lines_details = ""
    if hexagram_data['moving_lines']:
        if language == "it":
            moving_lines_details = "\n\n=== LINEE MUTEVOLI (CRUCIALI - SPIEGA OGNI LINEA IN DETTAGLIO) ===\n"
        else:
            moving_lines_details = "\n\n=== MOVING LINES (CRUCIAL - EXPLAIN EACH LINE IN DETAIL) ===\n"
        
        for line_pos in hexagram_data['moving_lines']:
            line_data = get_moving_line_extended(hexagram_data["primary_hexagram"], line_pos, language)
            testo = line_data.get("testo", "")
            significato = line_data.get("significato", "")
            if language == "it":
                moving_lines_details += f"\nLINEA {line_pos} MUTEVOLE:\nTesto tradizionale: \"{testo}\"\nSignificato: {significato}\n"
            else:
                moving_lines_details += f"\nMOVING LINE {line_pos}:\nTraditional text: \"{testo}\"\nMeaning: {significato}\n"
    
    # Derived hexagram details
    derived_details = ""
    if derived and derived_extended:
        derived_name = derived.get(name_key, derived.get("name", ""))
        derived_chinese = derived.get("name", "")
        derived_giudizio = derived_extended.get("giudizio", "")
        derived_immagine = derived_extended.get("immagine", "")
        
        if language == "it":
            derived_details = f"""

=== ESAGRAMMA DERIVATO (INDICA LA DIREZIONE FUTURA) ===
L'esagramma si trasforma in: {derived_chinese} - {derived_name}
Giudizio dell'esagramma derivato: "{derived_giudizio}"
Immagine: "{derived_immagine}"
Questo indica DOVE la situazione sta evolvendo e cosa aspettarsi nel futuro."""
        else:
            derived_details = f"""

=== DERIVED HEXAGRAM (INDICATES FUTURE DIRECTION) ===
The hexagram transforms into: {derived_chinese} - {derived_name}
Judgment of derived hexagram: "{derived_giudizio}"
Image: "{derived_immagine}"
This indicates WHERE the situation is evolving and what to expect in the future."""

    # Build conversation history context
    conversation_context = ""
    if conversation_history and len(conversation_history) > 0:
        if language == "it":
            conversation_context = "\n\n=== STORIA DELLA CONVERSAZIONE (Stese precedenti) ===\n"
            conversation_context += "Il consultante ha già fatto le seguenti domande in questa sessione. TIENI CONTO di questa storia per creare CONTINUITÀ nella risposta:\n"
            for i, prev in enumerate(conversation_history, 1):
                conversation_context += f"\n--- Stesa {i} ---\n"
                conversation_context += f"Domanda: \"{prev.get('question', '')}\"\n"
                conversation_context += f"Esagramma: {prev.get('hexagram_number')}. {prev.get('hexagram_name', '')}\n"
                if prev.get('derived_hexagram_number'):
                    conversation_context += f"Evolve in: {prev.get('derived_hexagram_number')}. {prev.get('derived_hexagram_name', '')}\n"
                # Include a summary of the previous interpretation (first 300 chars)
                prev_interp = prev.get('interpretation', '')[:300]
                conversation_context += f"Sintesi risposta: {prev_interp}...\n"
            conversation_context += "\nIMPORTANTE: La risposta attuale deve COLLEGARSI alle stese precedenti, creando una NARRAZIONE COERENTE. Fai riferimento a ciò che è emerso prima.\n"
        else:
            conversation_context = "\n\n=== CONVERSATION HISTORY (Previous readings) ===\n"
            conversation_context += "The querent has already asked the following questions in this session. TAKE THIS HISTORY INTO ACCOUNT to create CONTINUITY in your response:\n"
            for i, prev in enumerate(conversation_history, 1):
                conversation_context += f"\n--- Reading {i} ---\n"
                conversation_context += f"Question: \"{prev.get('question', '')}\"\n"
                conversation_context += f"Hexagram: {prev.get('hexagram_number')}. {prev.get('hexagram_name', '')}\n"
                if prev.get('derived_hexagram_number'):
                    conversation_context += f"Evolves into: {prev.get('derived_hexagram_number')}. {prev.get('derived_hexagram_name', '')}\n"
                prev_interp = prev.get('interpretation', '')[:300]
                conversation_context += f"Response summary: {prev_interp}...\n"
            conversation_context += "\nIMPORTANT: The current response must CONNECT to previous readings, creating a COHERENT NARRATIVE. Reference what emerged before.\n"

    # Build authoritative source context (Wilhelm Italian translation)
    wilhelm_context = build_authoritative_context(
        primary_number=hexagram_data["primary_hexagram"],
        derived_number=hexagram_data.get("derived_hexagram"),
        moving_lines=hexagram_data.get("moving_lines"),
        language="it",
    ) if language == "it" else ""

    if language == "it":
        user_prompt = f"""La domanda del consultante è: "{question}"
{conversation_context}
{wilhelm_context}

=== ESAGRAMMA PRINCIPALE ===
Nome: {primary_chinese} - {primary_name}
Numero: {hexagram_data["primary_hexagram"]}

TRIGRAMMA SUPERIORE: {trigramma_sup}
TRIGRAMMA INFERIORE: {trigramma_inf}

IL GIUDIZIO (SENTENZA TRADIZIONALE):
"{giudizio}"

L'IMMAGINE:
"{immagine}"

COMMENTO TRADIZIONALE:
{commento}
{moving_lines_details}{derived_details}

ISTRUZIONI:
Genera un'interpretazione RICCA, PROFONDA e DETTAGLIATA (600-900 parole) che:
1. Apra con una connessione poetica tra la domanda e il flusso del Tao
2. Spieghi in dettaglio il significato dell'esagramma e dei suoi trigrammi BASANDOTI sul testo Wilhelm sopra riportato
3. Citi LETTERALMENTE la Sentenza e l'Immagine dal testo originale Wilhelm e spiegale in relazione alla domanda specifica
4. SE CI SONO LINEE MUTEVOLI: dedica un paragrafo COMPLETO a ciascuna, riportando il TESTO TRADIZIONALE WILHELM («Nove al ...» o «Sei al ...») e analizzandolo per la situazione del consultante
5. SE C'È ESAGRAMMA DERIVATO: spiega la trasformazione usando il testo Wilhelm dell'esagramma derivato sopra, e cosa indica per il futuro
6. Concludi con saggezza pratica e un consiglio applicabile, ancorato al testo classico
{"7. SE C'È STORIA DELLA CONVERSAZIONE: collega questa risposta alle stese precedenti, creando una narrazione fluida" if conversation_context else ""}

CRITICO: Le tue interpretazioni devono essere FEDELI al testo Wilhelm sopra riportato.
Quando citi Sentenza, Immagine o linee mutevoli, USA LE PAROLE ESATTE di Wilhelm.
Scrivi come un antico maestro taoista, con poesia, profondità e compassione, ma sempre ANCORATO alla tradizione autentica.

==========================================================
DOPO l'interpretazione, aggiungi SEMPRE (anche se la domanda è breve) due
riassunti, esattamente in questo formato, con i marcatori in maiuscolo
identici a quelli qui sotto (sono parsati dal frontend, NON modificarli,
NON tradurli, NON aggiungere caratteri attorno):

===RIASSUNTO_RAPIDO===
[2-3 frasi, max 60 parole, che catturano l'essenza della risposta per il
consultante: cosa sta succedendo e cosa fare. Tono diretto e umano,
nessuna citazione classica.]
===RIASSUNTO_APPROFONDITO===
• Situazione attuale: [una frase ancorata all'esagramma principale]
• Insegnamento centrale: [una frase, citando in corsivo una parola-chiave
  della Sentenza Wilhelm]
• Linee mutevoli: [una frase sul significato cumulativo; "nessuna" se assenti]
• Azione consigliata: [una frase di consiglio applicabile oggi]
• Pericolo da evitare: [una frase chiara]
• Prospettiva: [una frase sull'esagramma derivato; "stabilità del primario"
  se non c'è derivato]
===FINE_RIASSUNTO==="""
    else:
        user_prompt = f"""The querent's question is: "{question}"
{conversation_context}
=== PRIMARY HEXAGRAM ===
Name: {primary_chinese} - {primary_name}
Number: {hexagram_data["primary_hexagram"]}

UPPER TRIGRAM: {trigramma_sup}
LOWER TRIGRAM: {trigramma_inf}

THE JUDGMENT (TRADITIONAL SENTENCE):
"{giudizio}"

THE IMAGE:
"{immagine}"

TRADITIONAL COMMENTARY:
{commento}
{moving_lines_details}{derived_details}

INSTRUCTIONS:
Generate a RICH, PROFOUND and DETAILED interpretation (600-900 words) that:
1. Opens with a poetic connection between the question and the flow of Tao
2. Explains in detail the meaning of the hexagram and its trigrams
3. Quotes and explains the Judgment and Image in relation to the specific question
4. IF THERE ARE MOVING LINES: dedicate a COMPLETE paragraph to each, explaining the traditional text and its meaning for the querent's situation
5. IF THERE IS A DERIVED HEXAGRAM: explain the transformation and what it indicates for the future
6. Conclude with practical wisdom and applicable advice

Write as an ancient Taoist master, with poetry, depth, and compassion.

==========================================================
AFTER the interpretation, ALWAYS append two summary blocks, in EXACTLY
this format. Keep the markers UPPERCASE and IDENTICAL to the ones below
— the frontend parses them, DO NOT translate or alter them:

===RIASSUNTO_RAPIDO===
[2-3 sentences, max 60 words, capturing the essence of the answer for the
querent: what is happening and what to do. Direct, human tone, no classical
citations.]
===RIASSUNTO_APPROFONDITO===
• Current situation: [one sentence anchored to the primary hexagram]
• Core teaching: [one sentence, italicizing one keyword from the Judgment]
• Moving lines: [one sentence on cumulative meaning; "none" if absent]
• Recommended action: [one sentence of applicable advice for today]
• Pitfall to avoid: [one clear sentence]
• Outlook: [one sentence on the derived hexagram; "stability of the primary"
  if there is no derived]
===FINE_RIASSUNTO==="""

    text = await _gemini_generate(
        system_instruction=system_prompt,
        prompt=user_prompt,
        generation_config=GEMINI_DEEP_CONFIG,
    )
    if text and len(text) > 100:
        return text
    logger.warning(f"Deep interp empty/short ({len(text or '')} char), using rich fallback")
    return _local_fallback_interpretation(
        primary, derived, primary_extended, derived_extended,
        hexagram_data, name_key, language, mode="deep",
    )

async def generate_direct_interpretation(hexagram_data: dict, question: str, language: str, 
                                          primary: dict, derived: dict, 
                                          primary_extended: dict, derived_extended: dict,
                                          name_key: str,
                                          conversation_history: list = None,
                                          topic: str = None) -> str:
    """Generate a direct, impactful interpretation - simple and to the point"""
    
    primary_name = primary.get(name_key, primary.get("name", ""))
    primary_chinese = primary.get("name", "")
    giudizio = primary_extended.get("giudizio", "")
    
    # Build moving lines summary
    moving_lines_text = ""
    if hexagram_data['moving_lines']:
        for line_pos in hexagram_data['moving_lines']:
            line_data = get_moving_line_extended(hexagram_data["primary_hexagram"], line_pos, language)
            testo = line_data.get("testo", "")
            if language == "it":
                moving_lines_text += f"\n• Linea {line_pos}: \"{testo}\""
            else:
                moving_lines_text += f"\n• Line {line_pos}: \"{testo}\""
    
    # Derived hexagram text
    derived_text = ""
    if derived:
        derived_name = derived.get(name_key, derived.get("name", ""))
        if language == "it":
            derived_text = f"\n\nLa situazione evolve verso: {derived_name}"
        else:
            derived_text = f"\n\nThe situation evolves towards: {derived_name}"
    
    # Build conversation history for direct style
    conversation_context = ""
    if conversation_history and len(conversation_history) > 0:
        if language == "it":
            conversation_context = "\n\nSTORIA DELLA CONVERSAZIONE:\n"
            for i, prev in enumerate(conversation_history, 1):
                conversation_context += f"- Domanda {i}: \"{prev.get('question', '')}\" → Esagramma {prev.get('hexagram_number')}\n"
            conversation_context += "\nCOLLEGA questa risposta alle precedenti in modo fluido.\n"
        else:
            conversation_context = "\n\nCONVERSATION HISTORY:\n"
            for i, prev in enumerate(conversation_history, 1):
                conversation_context += f"- Question {i}: \"{prev.get('question', '')}\" → Hexagram {prev.get('hexagram_number')}\n"
            conversation_context += "\nCONNECT this response to the previous ones fluidly.\n"

    # Topic context for focused interpretations
    topic_context_it = ""
    topic_context_en = ""
    if topic:
        topic_map_it = {
            'amore': 'AMORE E RELAZIONI - Interpreta tutto in chiave sentimentale e relazionale',
            'lavoro': 'LAVORO E CARRIERA - Interpreta tutto in chiave professionale e lavorativa',
            'fortuna': 'FORTUNA E OPPORTUNITÀ - Interpreta tutto in chiave di opportunità e destino',
            'soldi': 'FINANZE E DENARO - Interpreta tutto in chiave economica e finanziaria',
            'spirituale': 'CRESCITA SPIRITUALE - Interpreta tutto in chiave di evoluzione interiore',
            'personale': 'CRESCITA PERSONALE - Interpreta tutto in chiave di sviluppo personale'
        }
        topic_map_en = {
            'amore': 'LOVE AND RELATIONSHIPS - Interpret everything in romantic and relational terms',
            'lavoro': 'WORK AND CAREER - Interpret everything in professional and work terms',
            'fortuna': 'FORTUNE AND OPPORTUNITIES - Interpret everything in terms of opportunities and destiny',
            'soldi': 'FINANCES AND MONEY - Interpret everything in economic and financial terms',
            'spirituale': 'SPIRITUAL GROWTH - Interpret everything in terms of inner evolution',
            'personale': 'PERSONAL GROWTH - Interpret everything in terms of personal development'
        }
        topic_context_it = topic_map_it.get(topic, f'ARGOMENTO: {topic} - Interpreta tutto in relazione a questo tema specifico')
        topic_context_en = topic_map_en.get(topic, f'TOPIC: {topic} - Interpret everything in relation to this specific theme')

    if language == "it":
        topic_instruction = f"\n\n**ARGOMENTO SPECIFICO:** {topic_context_it}\nOgni parte della risposta DEVE essere focalizzata su questo argomento. Sii CONCRETO e SPECIFICO." if topic_context_it else ""
        
        system_prompt = f"""Sei un consulente I Ching che parla in modo DIRETTO, CHIARO e D'IMPATTO.
{topic_instruction}

STILE:
- Vai dritto al punto, senza giri di parole
- Usa un linguaggio semplice e comprensibile
- Parla SEMPRE in seconda persona ("tu", "la tua situazione")
- Sii empatico ma sincero - di' quello che il consultante ha bisogno di sentire
- Fornisci risposte pratiche e applicabili

STRUTTURA (300-400 parole):
1. Apertura diretta che conferma/risponde alla domanda (1-2 frasi d'impatto)
2. L'esagramma in sintesi: cosa significa per la situazione specifica
3. Se ci sono linee mutevoli: il messaggio chiave di ciascuna (una frase per linea)
4. Se c'è esagramma derivato: dove sta andando la situazione
5. Conclusione con consiglio pratico chiaro

NON FARE:
- Non usare linguaggio troppo poetico o elaborato
- Non fare lunghe citazioni
- Non essere vago o generico
- Non usare liste puntate (scrivi in paragrafi fluidi)

ESEMPIO DI TONO:
"La tua percezione è esatta. Quello che senti non è solo immaginazione - è reale. L'esagramma conferma che..."
"Ecco la verità sulla tua situazione: ..."
"Questo è il momento di..."
"""
    else:
        topic_instruction = f"\n\n**SPECIFIC TOPIC:** {topic_context_en}\nEvery part of the response MUST be focused on this topic. Be CONCRETE and SPECIFIC." if topic_context_en else ""
        
        system_prompt = f"""You are an I Ching consultant who speaks in a DIRECT, CLEAR and IMPACTFUL way.
{topic_instruction}

STYLE:
- Get straight to the point, no beating around the bush
- Use simple and understandable language
- ALWAYS speak in second person ("you", "your situation")
- Be empathetic but honest - say what the querent needs to hear
- Provide practical and applicable answers

STRUCTURE (300-400 words):
1. Direct opening that confirms/answers the question (1-2 impactful sentences)
2. The hexagram in summary: what it means for the specific situation
3. If there are moving lines: the key message of each (one sentence per line)
4. If there is a derived hexagram: where the situation is heading
5. Conclusion with clear practical advice

DO NOT:
- Do not use overly poetic or elaborate language
- Do not make long quotes
- Do not be vague or generic
- Do not use bullet lists (write in flowing paragraphs)

EXAMPLE TONE:
"Your perception is accurate. What you feel is not just imagination - it's real. The hexagram confirms that..."
"Here's the truth about your situation: ..."
"This is the time to..."
"""

    # Build authoritative Wilhelm context for direct mode (shorter excerpt)
    wilhelm_direct_context = ""
    if language == "it":
        from wilhelm_source import get_wilhelm_text
        primary_wilhelm = get_wilhelm_text(hexagram_data["primary_hexagram"], max_chars=2000)
        if primary_wilhelm:
            wilhelm_direct_context = f"""
=== TESTO ORIGINALE WILHELM (FONTE AUTORITATIVA) ===
{primary_wilhelm}
=== FINE TESTO WILHELM ==="""

    if language == "it":
        user_prompt = f"""Domanda del consultante: "{question}"
{conversation_context}
ESAGRAMMA: {hexagram_data["primary_hexagram"]}. {primary_chinese} ({primary_name})
Sentenza: "{giudizio}"
{f"Linee mutevoli: {moving_lines_text}" if moving_lines_text else "Nessuna linea mutevole"}
{derived_text}
{wilhelm_direct_context}

Genera un'interpretazione DIRETTA e D'IMPATTO (300-400 parole) che risponda chiaramente alla domanda.
USA il testo Wilhelm sopra come riferimento autorevole per il significato dell'esagramma e delle linee.
Quando rilevante, cita brevemente la Sentenza Wilhelm o la formula tradizionale di una linea mutevole.
Vai dritto al punto. Di' al consultante quello che ha bisogno di sapere — ancorato alla tradizione, non inventato.
{"Collega questa risposta alle domande precedenti nella conversazione." if conversation_context else ""}

DOPO l'interpretazione, aggiungi SEMPRE due riassunti con questi marcatori
ESATTI (in maiuscolo, sono parsati dal frontend, NON tradurli):

===RIASSUNTO_RAPIDO===
[1-2 frasi, max 40 parole: l'essenza della risposta. Nessuna citazione classica.]
===RIASSUNTO_APPROFONDITO===
• Situazione: [una frase ancorata all'esagramma]
• Insegnamento: [una frase chiave]
• Azione consigliata: [una frase concreta per oggi]
• Pericolo: [una frase chiara]
• Prospettiva: [una frase; "stabilità" se non c'è derivato]
===FINE_RIASSUNTO==="""
    else:
        user_prompt = f"""Querent's question: "{question}"
{conversation_context}
HEXAGRAM: {hexagram_data["primary_hexagram"]}. {primary_chinese} ({primary_name})
Judgment: "{giudizio}"
{f"Moving lines: {moving_lines_text}" if moving_lines_text else "No moving lines"}
{derived_text}

Generate a DIRECT and IMPACTFUL interpretation (300-400 words) that clearly answers the question.
Get straight to the point. Tell the querent what they need to know.
{"Connect this response to the previous questions in the conversation." if conversation_context else ""}

AFTER the interpretation, ALWAYS append two summary blocks with these
EXACT markers (uppercase, parsed by the frontend, DO NOT translate):

===RIASSUNTO_RAPIDO===
[1-2 sentences, max 40 words: the essence of the answer. No classical citations.]
===RIASSUNTO_APPROFONDITO===
• Situation: [one sentence anchored to the hexagram]
• Teaching: [one key sentence]
• Recommended action: [one concrete sentence for today]
• Pitfall: [one clear sentence]
• Outlook: [one sentence; "stability" if no derived]
===FINE_RIASSUNTO==="""

    text = await _gemini_generate(
        system_instruction=system_prompt,
        prompt=user_prompt,
        generation_config=GEMINI_DIRECT_CONFIG,
    )
    if text and len(text) > 100:
        return text
    logger.warning(f"Direct interp empty/short ({len(text or '')} char), using rich fallback")
    return _local_fallback_interpretation(
        primary, derived, primary_extended, derived_extended,
        hexagram_data, name_key, language, mode="direct",
    )

# ============== AUTH ROUTES ==============
@api_router.post("/auth/register", response_model=UserResponse)
@limiter.limit("10/hour")  # max 10 registrazioni/ora per IP → blocca creation spam
async def register(user_data: UserCreate, request: Request):
    # GDPR art. 6.1.a + art. 7: a valid legal basis requires the data subject
    # to have given consent BEFORE the data is processed. We refuse to create
    # the account if the privacy notice was not accepted.
    if not user_data.privacy_accepted:
        raise HTTPException(
            status_code=400,
            detail="È necessario accettare l'Informativa Privacy per registrarsi."
        )

    existing = await db.users.find_one({"email": user_data.email})
    if existing:
        raise HTTPException(status_code=400, detail="Email già registrata")

    # Record proof of consent (art. 7.1 GDPR — controller must demonstrate consent)
    now_iso = datetime.now(timezone.utc).isoformat()
    client_ip = request.client.host if request and request.client else None
    user_agent = request.headers.get("User-Agent", "") if request else ""

    user_id = str(uuid.uuid4())
    user_doc = {
        "id": user_id,
        "email": user_data.email,
        "password": hash_password(user_data.password),
        "name": user_data.name,
        "phone": user_data.phone,
        "language": user_data.language,
        "subscription_active": False,
        "subscription_end": None,
        "created_at": now_iso,
        "consents": {
            "privacy_accepted": True,
            "privacy_accepted_at": now_iso,
            "privacy_version": "2.0",  # bump when policy text changes
            "marketing_consent": bool(user_data.marketing_consent),
            "marketing_consent_at": now_iso if user_data.marketing_consent else None,
            "ip": client_ip,
            "user_agent": user_agent[:300],
        },
    }
    await db.users.insert_one(user_doc)

    # Email di benvenuto — no-op se Resend non configurato. Non blocca
    # la registrazione anche se l'invio fallisce.
    try:
        mailer.send_welcome(to=user_data.email, user_name=user_data.name or "")
    except Exception as e:  # noqa: BLE001
        logger.warning("Welcome email failed for %s: %s", user_data.email, e)

    return UserResponse(
        id=user_id,
        email=user_data.email,
        name=user_data.name,
        phone=user_data.phone,
        language=user_data.language,
        subscription_active=False
    )

@api_router.post("/auth/login")
@limiter.limit("10/minute")  # max 10 login/min per IP → brute-force ridotto
async def login(request: Request, credentials: UserLogin):
    # Lockout per email: 5 fallimenti consecutivi entro 15 minuti -> blocco
    # 30 minuti, indipendente dall'IP (un attaccante che ruota IP residuali
    # non sfugge al lockout perche' lo applichiamo per identita').
    now = datetime.now(timezone.utc)
    lockout_window = timedelta(minutes=15)
    lockout_duration = timedelta(minutes=30)
    user = await db.users.find_one({"email": credentials.email}, {"_id": 0})
    if user:
        fails = user.get("failed_login_attempts") or []
        # Mantieni solo i timestamp recenti
        recent_fails = []
        for ts in fails:
            try:
                t = datetime.fromisoformat(str(ts).replace("Z", "+00:00"))
                if now - t < lockout_window:
                    recent_fails.append(t)
            except Exception:
                continue
        # Se siamo in lockout, rifiuta direttamente
        locked_until = user.get("locked_until")
        if locked_until:
            try:
                lu = datetime.fromisoformat(str(locked_until).replace("Z", "+00:00"))
                if lu > now:
                    raise HTTPException(
                        status_code=429,
                        detail=f"Troppi tentativi falliti. Riprova tra {int((lu - now).total_seconds() / 60) + 1} minuti."
                    )
            except HTTPException:
                raise
            except Exception:
                pass

    if not user or not verify_password(credentials.password, user["password"]):
        # Registra il fallimento. Se l'utente NON esiste non lo facciamo (non
        # vogliamo creare record fantasma per enumerazione).
        if user:
            fails = (user.get("failed_login_attempts") or [])[-9:]  # cap 10
            fails.append(now.isoformat())
            update = {"failed_login_attempts": fails}
            if len([t for t in recent_fails] + [now]) >= 5:
                update["locked_until"] = (now + lockout_duration).isoformat()
            await db.users.update_one({"id": user["id"]}, {"$set": update})
        raise HTTPException(status_code=401, detail="Credenziali non valide")

    # Login OK: pulisci storico fallimenti
    if user.get("failed_login_attempts") or user.get("locked_until"):
        await db.users.update_one(
            {"id": user["id"]},
            {"$set": {"failed_login_attempts": [], "locked_until": None}},
        )

    token = create_token(user["id"], user["email"])

    admin = is_admin_email(user.get("email"))
    plan = get_user_plan(user)
    display = user.get("display_name") or user.get("name", "")
    return {
        "token": token,
        "user": UserResponse(
            id=user["id"],
            email=user["email"],
            name=display,
            phone=user.get("phone", ""),
            language=user["language"],
            subscription_active=user.get("subscription_active", False) or admin,
            subscription_end=user.get("subscription_end"),
            is_admin=admin,
            plan=plan,
        )
    }

GOOGLE_CLIENT_ID = os.environ.get("GOOGLE_CLIENT_ID", "")


@api_router.post("/auth/google")
async def google_id_token_login(data: dict):
    """
    Login/Register via Google Identity Services (modern OAuth flow).

    The frontend Google button produces a signed JWT (id_token / 'credential').
    We verify it server-side against Google's public keys, then create
    or update the user and return our own session JWT.

    Body: { "credential": "<google_id_token_jwt>" }
    """
    if not GOOGLE_CLIENT_ID:
        raise HTTPException(
            status_code=501,
            detail="Google Sign-In non configurato. Imposta GOOGLE_CLIENT_ID."
        )

    credential = data.get("credential") or data.get("id_token")
    if not credential:
        raise HTTPException(status_code=400, detail="credential (id_token) richiesto")

    try:
        from google.oauth2 import id_token as google_id_token
        from google.auth.transport import requests as google_requests
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="google-auth non installato sul backend"
        )

    # Verify the token signature, expiry, and audience (must match our Client ID)
    try:
        idinfo = google_id_token.verify_oauth2_token(
            credential,
            google_requests.Request(),
            GOOGLE_CLIENT_ID,
            clock_skew_in_seconds=10,
        )
    except ValueError as e:
        logger.warning(f"Google token verification failed: {e}")
        raise HTTPException(status_code=401, detail="Token Google non valido o scaduto")

    # Extract identity from verified payload
    google_email = idinfo.get("email")
    google_name = idinfo.get("name", "")
    google_picture = idinfo.get("picture", "")
    google_email_verified = idinfo.get("email_verified", False)
    google_sub = idinfo.get("sub")  # Stable Google user id

    if not google_email or not google_email_verified:
        raise HTTPException(status_code=400, detail="Email Google non verificata")

    # Find or create user
    existing_user = await db.users.find_one({"email": google_email}, {"_id": 0})

    if existing_user:
        user_id = existing_user["id"]
        user_name = existing_user.get("name") or google_name
        user_language = existing_user.get("language", "it")
        subscription_active = existing_user.get("subscription_active", False)
        subscription_end = existing_user.get("subscription_end")
        # Link Google fields if missing
        await db.users.update_one(
            {"id": user_id},
            {"$set": {
                "google_sub": google_sub,
                "google_picture": google_picture,
                "google_name": google_name,
                "last_login": datetime.now(timezone.utc).isoformat(),
            }}
        )
    else:
        # First-time Google user → create account (free plan)
        user_id = str(uuid.uuid4())
        user_doc = {
            "id": user_id,
            "email": google_email,
            "password": "",  # No password for OAuth-only users
            "name": google_name,
            "phone": "",
            "language": "it",
            "google_sub": google_sub,
            "google_picture": google_picture,
            "google_name": google_name,
            "subscription_active": False,
            "subscription_end": None,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "auth_provider": "google",
        }
        await db.users.insert_one(user_doc)
        user_name = google_name
        user_language = "it"
        subscription_active = False
        subscription_end = None

    # Issue our own JWT
    token = create_token(user_id, google_email)

    # Admin whitelist (owner) gets Premium for free regardless of payments
    admin = is_admin_email(google_email)
    plan = "premium" if (admin or subscription_active) else "free"

    return {
        "token": token,
        "user": {
            "id": user_id,
            "email": google_email,
            "name": user_name,
            "picture": google_picture,
            "language": user_language,
            "subscription_active": bool(subscription_active or admin),
            "subscription_end": subscription_end,
            "is_admin": admin,
            "plan": plan,
        }
    }


# Backwards-compat alias for any old client still hitting /callback
@api_router.post("/auth/google/callback")
async def google_oauth_callback_compat(data: dict):
    return await google_id_token_login(data)

@api_router.get("/auth/me", response_model=UserResponse)
async def get_me(user: dict = Depends(get_current_user)):
    # If the user has activated a pseudonym, display that instead of the real name.
    # Real name stays in the DB so they can revert later (see /auth/pseudonym).
    display = user.get("display_name") or user.get("name", "")
    admin = is_admin_email(user.get("email"))
    plan = get_user_plan(user)  # 'premium' for admins or paying users
    return UserResponse(
        id=user["id"],
        email=user["email"],
        name=display,
        phone=user.get("phone", ""),
        language=user["language"],
        subscription_active=user.get("subscription_active", False) or admin,
        subscription_end=user.get("subscription_end"),
        is_admin=admin,
        plan=plan,
    )

@api_router.put("/auth/language")
async def update_language(language: str, user: dict = Depends(get_current_user)):
    if language not in ["it", "en"]:
        raise HTTPException(status_code=400, detail="Lingua non supportata")
    await db.users.update_one({"id": user["id"]}, {"$set": {"language": language}})
    return {"message": "Lingua aggiornata"}


@api_router.put("/auth/pseudonym")
async def set_pseudonym(data: dict, user: dict = Depends(get_current_user)):
    """
    Honors the Privacy Policy promise:
    "Su richiesta puoi attivare l'uso di uno pseudonimo al posto del nome reale."

    Accepts:
      { "pseudonym": "<string>" }  -> sets it (displayed in UI in place of real name)
      { "pseudonym": null }        -> removes it (real name resumes)

    The real name is preserved in the DB (so the user can revert) but the
    public-facing `name` is overridden by `display_name` everywhere the UI
    reads it from /auth/me, /profile, /history shared cards, etc.
    """
    pseudonym = data.get("pseudonym")
    if pseudonym is not None:
        pseudonym = str(pseudonym).strip()
        if pseudonym == "":
            pseudonym = None
        elif len(pseudonym) > 40:
            raise HTTPException(status_code=400, detail="Lo pseudonimo non può superare 40 caratteri")

    update = {
        "display_name": pseudonym,
        "pseudonym_active": pseudonym is not None,
        "pseudonym_updated_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.users.update_one({"id": user["id"]}, {"$set": update})
    return {"message": "Pseudonimo aggiornato", "pseudonym_active": pseudonym is not None, "display_name": pseudonym}


@api_router.put("/auth/marketing-consent")
async def update_marketing_consent(data: dict, user: dict = Depends(get_current_user)):
    """
    Allow user to grant or revoke marketing consent at any time (art. 7.3 GDPR).
    Body: { "consent": true | false }
    """
    consent = bool(data.get("consent", False))
    now = datetime.now(timezone.utc).isoformat()
    await db.users.update_one(
        {"id": user["id"]},
        {"$set": {
            "consents.marketing_consent": consent,
            "consents.marketing_consent_at": now if consent else None,
            "consents.marketing_revoked_at": None if consent else now,
        }}
    )
    return {"message": "Preferenze marketing aggiornate", "marketing_consent": consent}


@api_router.delete("/auth/account")
async def delete_account(user: dict = Depends(get_current_user)):
    """
    GDPR art. 17 — Right to erasure ("right to be forgotten").

    Deletes the user's account and all related personal data. Consultations
    are anonymized (user_id removed, retained in aggregate-statistics form)
    rather than fully deleted, as allowed by recital 26 for purely statistical
    purposes that don't allow re-identification.
    """
    uid = user["id"]
    email = user.get("email", "")

    # Anonymize past consultations (keep only the I Ching content, drop link to user)
    await db.consultations.update_many(
        {"user_id": uid},
        {"$set": {"user_id": None, "anonymized": True, "anonymized_at": datetime.now(timezone.utc).isoformat()},
         "$unset": {"question": ""}}  # question may contain personal info -> drop
    )

    # Delete user-private collections completely (per user_id)
    # Aggiunte le collezioni che il vecchio codice DIMENTICAVA e che sono
    # emerse dall'audit: fitness_*, notifications, notification_reads,
    # coin_toss_sessions, badges assegnati.
    for coll in (
        "notes",
        "user_paths",
        "notification_preferences",
        "payment_transactions",
        "completed_paths",
        "fitness_profiles",
        "fitness_programs",
        "fitness_completed",
        "notifications",
        "notification_reads",
        "coin_toss_sessions",
        "user_badges",
        "user_level_progress",
    ):
        try:
            await db[coll].delete_many({"user_id": uid})
        except Exception as e:  # noqa: BLE001
            logger.warning("delete_account: skip collection %s: %s", coll, e)

    # Collezioni indicizzate per email (non user_id): password reset requests
    # e refund requests. Le eliminiamo esplicitamente.
    if email:
        for coll in ("password_resets", "refund_requests"):
            try:
                await db[coll].delete_many({"email": email})
            except Exception as e:  # noqa: BLE001
                logger.warning("delete_account: skip email-coll %s: %s", coll, e)

    # Finally remove the user record itself
    await db.users.delete_one({"id": uid})

    # Audit trail — teniamo solo un record minimo per adempiere all'obbligo
    # di dimostrare l'avvenuta cancellazione (GDPR art. 5.2 accountability).
    # Zero PII: solo hash SHA-256 dell'email + timestamp.
    import hashlib
    email_hash = hashlib.sha256((email or "").lower().encode()).hexdigest()[:16] if email else None
    try:
        await db.deletion_log.insert_one({
            "email_hash": email_hash,
            "deleted_at": datetime.now(timezone.utc).isoformat(),
        })
    except Exception:
        pass  # log opzionale, non blocchiamo

    return {"message": "Account eliminato. Tutti i dati personali sono stati cancellati."}


@api_router.get("/auth/export")
async def export_user_data(user: dict = Depends(get_current_user)):
    """
    GDPR art. 20 — Right to data portability.
    Returns a JSON dump of all data the user can claim ownership of.
    """
    uid = user["id"]
    user_doc = await db.users.find_one({"id": uid}, {"_id": 0, "password": 0})
    consultations = await db.consultations.find({"user_id": uid}, {"_id": 0}).to_list(10000)
    notes = await db.notes.find({"user_id": uid}, {"_id": 0}).to_list(10000)
    paths = await db.user_paths.find({"user_id": uid}, {"_id": 0}).to_list(1000)
    prefs = await db.notification_preferences.find_one({"user_id": uid}, {"_id": 0})
    return {
        "exported_at": datetime.now(timezone.utc).isoformat(),
        "format_version": "1.0",
        "user": user_doc,
        "consultations": consultations,
        "notes": notes,
        "paths": paths,
        "notification_preferences": prefs,
    }

@api_router.post("/auth/request-reset")
@limiter.limit("5/hour")  # max 5 richieste reset/ora per IP → niente spam email
async def request_password_reset(request: Request, data: PasswordResetRequest):
    """
    Richiede il reset della password.
    Genera un codice temporaneo, lo salva su DB e lo invia via email.
    Risposta SEMPRE identica per evitare user enumeration.
    """
    # Il messaggio finale è IDENTICO per utente esistente e non esistente
    # (prima erano diversi -> permetteva enumeration via response body length).
    GENERIC_MSG = (
        "Se l'email risulta registrata, riceverai a breve un'email "
        "con un codice di 8 cifre per reimpostare la password. "
        "Controlla anche la cartella spam."
    )

    user = await db.users.find_one({"email": data.email}, {"_id": 0})
    if not user:
        # Aggiungiamo delay pseudo-casuale per equalizzare i tempi con il
        # ramo "utente esiste" che fa insert Mongo + chiamata Resend
        # (~150-400ms). Senza questo, un attaccante puo' misurare i tempi
        # per capire quali email esistono.
        import asyncio, random as _rnd
        await asyncio.sleep(_rnd.uniform(0.15, 0.4))
        return {"message": GENERIC_MSG}
    
    # Genera codice di reset
    reset_code = generate_reset_code()
    expires_at = datetime.now(timezone.utc) + timedelta(hours=1)
    
    # Salva la richiesta di reset nel database
    reset_request = {
        "id": str(uuid.uuid4()),
        "user_id": user["id"],
        "email": data.email,
        "phone": data.phone or user.get("phone", ""),
        "user_name": user["name"],
        "code": reset_code,
        "expires_at": expires_at.isoformat(),
        "used": False,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.password_resets.insert_one(reset_request)

    # ────────── INVIO EMAIL VERA via Resend ──────────
    # Se Resend non e' configurato (RESEND_API_KEY mancante), il codice viene
    # comunque salvato e loggato per l'admin — l'utente vede solo il messaggio
    # generico, MAI il codice in chiaro nella risposta JSON (era un leak).
    email_sent = mailer.send_password_reset(
        to=data.email,
        user_name=user.get("name") or "",
        code=reset_code,
        expires_minutes=60,
    )
    if not email_sent:
        # Resend off o errore: log per l'admin con codice (visibile solo nei log Render)
        logger.warning(
            "🔐 RESET PASSWORD (email NON inviata, recupero manuale richiesto) "
            "email=%s nome=%s codice=%s scade=%s",
            data.email, user.get("name"), reset_code, expires_at.isoformat()
        )

    # Stessa stringa identica al ramo "utente non esiste" (definita in cima
    # alla funzione come GENERIC_MSG) per non trapelare informazioni.
    return {"message": GENERIC_MSG}

@api_router.post("/auth/verify-reset")
@limiter.limit("10/hour")  # max 10 tentativi/ora per IP → ferma bruteforce codice
async def verify_reset_code(request: Request, data: PasswordResetVerify):
    """
    Verifica il codice di reset e imposta la nuova password.
    Confronto del codice in tempo costante (hmac.compare_digest) per non
    aprire un timing channel su codici parzialmente indovinati.
    """
    # 1. Recupera l'unico record NON usato piu' recente per questa email
    reset_request = await db.password_resets.find_one(
        {"email": data.email, "used": False},
        {"_id": 0},
        sort=[("created_at", -1)],
    )
    if not reset_request:
        raise HTTPException(status_code=400, detail="Codice non valido o già utilizzato")

    # 2. Confronto costante del codice (previene timing side-channel).
    #    hmac.compare_digest richiede stringhe della stessa lunghezza:
    #    normalizziamo a str e usiamo il segno hmac su bytes.
    stored_code = str(reset_request.get("code", ""))
    supplied_code = str(data.code or "")
    if not hmac.compare_digest(stored_code, supplied_code):
        raise HTTPException(status_code=400, detail="Codice non valido o già utilizzato")
    
    # Verifica scadenza
    expires_at = datetime.fromisoformat(reset_request["expires_at"].replace("Z", "+00:00"))
    if expires_at < datetime.now(timezone.utc):
        raise HTTPException(status_code=400, detail="Codice scaduto. Richiedi un nuovo reset.")
    
    # Valida la nuova password
    if len(data.new_password) < 6:
        raise HTTPException(status_code=400, detail="La password deve essere almeno 6 caratteri")
    
    # Aggiorna la password dell'utente
    hashed_password = hash_password(data.new_password)
    await db.users.update_one(
        {"email": data.email},
        {"$set": {"password": hashed_password}}
    )
    
    # Segna il codice come usato
    await db.password_resets.update_one(
        {"id": reset_request["id"]},
        {"$set": {"used": True, "used_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    logger.info(f"✅ Password resettata per: {data.email}")
    
    return {"message": "Password aggiornata con successo. Ora puoi accedere."}

# ═══════════════════════════════════════════════════════════════════════
# NOTIFICHE INTELLIGENTI in-app
# ═══════════════════════════════════════════════════════════════════════
#
# Ogni notifica è generata server-side a partire da:
#   - esagramma I Ching del giorno (uguale per tutti)
#   - fase lunare attuale
#   - energia del calendario cinese (Jiazi cycle)
#   - eventuali percorsi guidati attivi dell'utente
#   - aggiornamenti del programma Fitness (se piano fitness_coaching)
#
# Il client le scarica con GET /api/notifications/inbox e le marca
# come lette con POST /api/notifications/{id}/read.

@api_router.get("/notifications/inbox")
async def get_notifications_inbox(request: Request, user: dict = Depends(get_current_user)):
    """
    Returns the user's intelligent notification feed.
    Mixes:
      - automatic system notifications (lunar, daily I Ching, etc.)
      - completed-paths unread counter
      - stored notifications (db.notifications) if any
    """
    lang = user.get("language", "it")
    notifications = []

    # 1) Notifica esagramma del giorno
    today = datetime.now(timezone.utc).date()
    hex_n = get_daily_hexagram_number()
    hex_data = HEXAGRAMS.get(hex_n, {}) or {}
    name_key = "name_it" if lang == "it" else "name_en"
    hex_name = hex_data.get(name_key, hex_data.get("name", ""))
    notifications.append({
        "id": f"daily_hex_{today.isoformat()}",
        "type": "daily_hexagram",
        "icon": "📖",
        "title": (f"Esagramma del giorno: {hex_name}" if lang == "it"
                  else f"Today's hexagram: {hex_name}"),
        "body": (f"Oggi il Tao ti parla attraverso #{hex_n} {hex_name}. Lascia che la sua energia ti guidi."
                 if lang == "it"
                 else f"Today the Tao speaks through #{hex_n} {hex_name}. Let its energy guide you."),
        "deeplink": "/library/" + str(hex_n),
        "created_at": datetime.combine(today, datetime.min.time(), tzinfo=timezone.utc).isoformat(),
        "read": False,
    })

    # 2) Notifica fase lunare
    lunar = get_lunar_phase()
    notifications.append({
        "id": f"lunar_{today.isoformat()}",
        "type": "lunar_phase",
        "icon": lunar.get("emoji", "🌙"),
        "title": (lunar.get("name_it", "") if lang == "it" else lunar.get("name_en", "")),
        "body": (lunar.get("advice_it", "") if lang == "it" else lunar.get("advice_en", "")),
        "deeplink": "/library",
        "created_at": datetime.combine(today, datetime.min.time(), tzinfo=timezone.utc).isoformat(),
        "read": False,
    })

    # 3) Calendario cinese — energia del giorno
    try:
        day_energy = get_chinese_day_energy()
        notifications.append({
            "id": f"chinese_{today.isoformat()}",
            "type": "chinese_calendar",
            "icon": day_energy.get("animal", {}).get("emoji", "🐉"),
            "title": (f"Energia di oggi: {day_energy.get('element')}" if lang == "it"
                      else f"Today's energy: {day_energy.get('element_en')}"),
            "body": (day_energy.get("quality_it", "") if lang == "it" else day_energy.get("quality_en", "")),
            "deeplink": "/dashboard",
            "created_at": datetime.combine(today, datetime.min.time(), tzinfo=timezone.utc).isoformat(),
            "read": False,
        })
    except Exception:
        pass

    # 4) Percorsi completati non letti
    try:
        unread_paths = await db.completed_paths.count_documents({
            "user_id": user["id"],
            "is_read": False,
        })
        if unread_paths > 0:
            notifications.append({
                "id": "unread_paths",
                "type": "completed_paths",
                "icon": "🛤️",
                "title": (f"{unread_paths} percorso/i completato/i da rivedere" if lang == "it"
                          else f"{unread_paths} completed path(s) to review"),
                "body": (f"Hai {unread_paths} sintesi di percorso pronte da leggere."
                         if lang == "it"
                         else f"You have {unread_paths} path syntheses ready to read."),
                "deeplink": "/completed-paths",
                "created_at": datetime.now(timezone.utc).isoformat(),
                "read": False,
            })
    except Exception:
        pass

    # 5) Programma Fitness — promemoria attività di oggi
    plan = get_user_plan(user)
    if get_plan_limits(plan).get("can_fitness_coaching"):
        try:
            program = await db.fitness_programs.find_one(
                {"user_id": user["id"], "active": True}, {"_id": 0}
            )
            if program:
                today_str = today.isoformat()
                for day in program.get("days", []):
                    if day.get("date") == today_str:
                        pending = [a for a in day.get("activities", []) if not a.get("completed")]
                        if pending:
                            notifications.append({
                                "id": f"fitness_today_{today_str}",
                                "type": "fitness_reminder",
                                "icon": "✨",
                                "title": (f"{len(pending)} attività Fitness in programma oggi"
                                          if lang == "it"
                                          else f"{len(pending)} fitness activities scheduled today"),
                                "body": ", ".join(a.get("title", "")[:35] for a in pending[:3]),
                                "deeplink": "/fitness",
                                "created_at": datetime.now(timezone.utc).isoformat(),
                                "read": False,
                            })
                        break
        except Exception:
            pass

    # 6) Mark which ones the user has already read (stored read state)
    try:
        read_doc = await db.notification_reads.find_one(
            {"user_id": user["id"]}, {"_id": 0, "read_ids": 1}
        )
        read_ids = set((read_doc or {}).get("read_ids", []))
        for n in notifications:
            if n["id"] in read_ids:
                n["read"] = True
    except Exception:
        pass

    unread_count = sum(1 for n in notifications if not n["read"])
    return {
        "notifications": notifications,
        "unread_count": unread_count,
        "total": len(notifications),
    }


@api_router.post("/notifications/{notif_id}/read")
async def mark_notification_read(notif_id: str, user: dict = Depends(get_current_user)):
    """Persist that this notification was acknowledged by the user."""
    await db.notification_reads.update_one(
        {"user_id": user["id"]},
        {"$addToSet": {"read_ids": notif_id}, "$set": {"updated_at": datetime.now(timezone.utc).isoformat()}},
        upsert=True,
    )
    return {"id": notif_id, "read": True}


@api_router.post("/notifications/mark-all-read")
async def mark_all_notifications_read(user: dict = Depends(get_current_user)):
    """Mark every currently-fetchable notification as read."""
    box = await get_notifications_inbox.__wrapped__(None, user) if hasattr(get_notifications_inbox, "__wrapped__") else None
    # We can't easily reuse the FastAPI handler; just clear by inserting today's keys
    today = datetime.now(timezone.utc).date().isoformat()
    keys = [
        f"daily_hex_{today}",
        f"lunar_{today}",
        f"chinese_{today}",
        "unread_paths",
        f"fitness_today_{today}",
    ]
    await db.notification_reads.update_one(
        {"user_id": user["id"]},
        {"$addToSet": {"read_ids": {"$each": keys}},
         "$set": {"updated_at": datetime.now(timezone.utc).isoformat()}},
        upsert=True,
    )
    return {"marked": len(keys)}


# ═══════════════════════════════════════════════════════════════════════
# FITNESS COACHING (esclusivo piano fitness_coaching)
# ═══════════════════════════════════════════════════════════════════════

def _require_fitness_plan(user: dict):
    """Raises 403 unless the user's plan unlocks Fitness Coaching."""
    plan = get_user_plan(user)
    limits = get_plan_limits(plan)
    if not limits.get("can_fitness_coaching"):
        raise HTTPException(
            status_code=403,
            detail="Il programma Benessere Fisico è disponibile solo nel piano dedicato."
        )


@api_router.get("/fitness/onboarding/questions")
async def fitness_get_onboarding(user: dict = Depends(get_current_user)):
    """Returns the onboarding questionnaire. Open to any logged-in user
    so the free preview can show what the program looks like."""
    return {"questions": FITNESS_ONBOARDING_QUESTIONS}


@api_router.get("/fitness/onboarding")
async def fitness_get_my_onboarding(user: dict = Depends(get_current_user)):
    """Read back the user's saved answers (if any)."""
    profile = await db.fitness_profiles.find_one({"user_id": user["id"]}, {"_id": 0})
    return {
        "has_profile": bool(profile),
        "profile": profile,
        "focus": fitness_score(profile.get("answers", {})) if profile else None,
    }


@api_router.post("/fitness/onboarding")
async def fitness_save_onboarding(data: dict, user: dict = Depends(get_current_user)):
    """Save the answers. Required before generating a program."""
    _require_fitness_plan(user)
    answers = data.get("answers") or {}
    validation = fitness_validate_onboarding(answers)
    if not validation["ok"]:
        raise HTTPException(status_code=400, detail={
            "message": "Onboarding incompleto",
            "missing": validation["missing"],
            "errors": validation["errors"],
        })

    focus = fitness_score(answers)
    doc = {
        "user_id": user["id"],
        "answers": answers,
        "focus": focus,
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.fitness_profiles.update_one(
        {"user_id": user["id"]},
        {"$set": doc, "$setOnInsert": {"created_at": doc["updated_at"]}},
        upsert=True,
    )
    return {"saved": True, "focus": focus}


@api_router.post("/fitness/program/generate")
async def fitness_generate_new_program(user: dict = Depends(get_current_user)):
    """Generate a fresh weekly program based on the saved onboarding."""
    _require_fitness_plan(user)

    profile = await db.fitness_profiles.find_one({"user_id": user["id"]}, {"_id": 0})
    if not profile or not profile.get("answers"):
        raise HTTPException(
            status_code=400,
            detail="Completa prima il questionario di onboarding."
        )

    # Use today's I Ching hexagram as a thematic anchor
    hex_number = get_daily_hexagram_number()
    hex_data = HEXAGRAMS.get(hex_number, {}) or {}
    lang = user.get("language", "it")
    name_key = "name_it" if lang == "it" else "name_en"
    hex_name = hex_data.get(name_key, hex_data.get("name", ""))

    program = fitness_generate_program(
        user_id=user["id"],
        answers=profile["answers"],
        iching_hexagram_number=hex_number,
        iching_hexagram_name=hex_name,
        language=lang,
    )

    # Replace any current program with this new one (one active at a time)
    await db.fitness_programs.delete_many({"user_id": user["id"], "active": True})
    program_record = {**program, "active": True}
    await db.fitness_programs.insert_one(program_record)

    # Return a clean copy (strip Mongo's _id)
    program_record.pop("_id", None)
    return program_record


@api_router.get("/fitness/program/current")
async def fitness_get_current_program(user: dict = Depends(get_current_user)):
    """Return the user's currently active program, if any."""
    _require_fitness_plan(user)
    program = await db.fitness_programs.find_one(
        {"user_id": user["id"], "active": True},
        {"_id": 0},
    )
    if not program:
        return {"has_program": False, "program": None}
    return {"has_program": True, "program": program}


@api_router.post("/fitness/activity/{activity_id}/complete")
async def fitness_complete_activity(
    activity_id: str,
    user: dict = Depends(get_current_user),
):
    """Mark a single activity as completed (or toggle it back)."""
    _require_fitness_plan(user)
    program = await db.fitness_programs.find_one(
        {"user_id": user["id"], "active": True},
        {"_id": 0},
    )
    if not program:
        raise HTTPException(status_code=404, detail="Nessun programma attivo.")

    changed = False
    completed_now = False
    for day in program.get("days", []):
        for act in day.get("activities", []):
            if act.get("id") == activity_id:
                act["completed"] = not act.get("completed", False)
                completed_now = act["completed"]
                if completed_now:
                    act["completed_at"] = datetime.now(timezone.utc).isoformat()
                else:
                    act.pop("completed_at", None)
                changed = True
                break
        if changed:
            break

    if not changed:
        raise HTTPException(status_code=404, detail="Attività non trovata nel programma.")

    await db.fitness_programs.update_one(
        {"user_id": user["id"], "active": True},
        {"$set": {"days": program["days"], "updated_at": datetime.now(timezone.utc).isoformat()}},
    )

    # If marking complete, record in completed_activities for cross-program stats
    if completed_now:
        await db.fitness_completed.insert_one({
            "user_id": user["id"],
            "activity_id": activity_id,
            "program_id": program["id"],
            "completed_at": datetime.now(timezone.utc).isoformat(),
        })

    return {"activity_id": activity_id, "completed": completed_now}


@api_router.get("/fitness/stats")
async def fitness_get_stats(user: dict = Depends(get_current_user)):
    """Aggregated XP / streak / badges across all programs."""
    _require_fitness_plan(user)
    total_done = await db.fitness_completed.count_documents({"user_id": user["id"]})

    # Streak: consecutive days with at least one completion (work backwards)
    today = datetime.now(timezone.utc).date()
    streak = 0
    for i in range(30):  # cap at 30 days back
        d = today - timedelta(days=i)
        start = datetime(d.year, d.month, d.day, tzinfo=timezone.utc)
        end = start + timedelta(days=1)
        count = await db.fitness_completed.count_documents({
            "user_id": user["id"],
            "completed_at": {"$gte": start.isoformat(), "$lt": end.isoformat()},
        })
        if count == 0:
            if i == 0:
                continue  # today might still be empty
            break
        streak += 1

    return fitness_compute_xp(total_done, streak)


# ═══════════════════════════════════════════════════════════════════════
# ADMIN
# ═══════════════════════════════════════════════════════════════════════
@api_router.get("/admin/reset-requests")
async def get_reset_requests(request: Request):
    """
    Endpoint admin per vedere le richieste di reset pendenti.

    AUTH: X-Admin-Secret header DEVE corrispondere a ADMIN_SECRET env var.
    Prima dell'aggiunta di questa verifica l'endpoint era pubblico ed
    esponeva codici reset password in chiaro — CVE interna.
    """
    _verify_admin(request.headers.get("X-Admin-Secret"))
    requests = await db.password_resets.find(
        {"used": False},
        {"_id": 0}
    ).sort("created_at", -1).to_list(50)
    return requests


def _verify_admin(secret_header: str):
    """Constant-time check against ADMIN_SECRET env var."""
    import hmac
    expected = os.environ.get("ADMIN_SECRET", "")
    if not expected:
        raise HTTPException(status_code=503, detail="ADMIN_SECRET non configurato")
    if not hmac.compare_digest(secret_header or "", expected):
        raise HTTPException(status_code=403, detail="Accesso admin non autorizzato")


@api_router.post("/admin/revoke-test-premium")
async def revoke_test_premium(request: Request):
    """
    Revokes Premium from any user whose subscription was activated manually
    (i.e. they have subscription_active=true but no associated paid
    payment_transaction). End of the testing window.

    Auth: X-Admin-Secret header must match the ADMIN_SECRET env var.
    """
    _verify_admin(request.headers.get("X-Admin-Secret"))

    # Find paid transactions to know which users LEGITIMATELY have Premium
    paid = await db.payment_transactions.distinct("user_id", {"payment_status": "paid"})
    paid_set = set(paid)

    affected = []
    kept_admins = []
    premium_users = await db.users.find(
        {"subscription_active": True},
        {"_id": 0, "id": 1, "email": 1, "subscription_end": 1, "auth_provider": 1}
    ).to_list(1000)

    for u in premium_users:
        if u["id"] in paid_set:
            continue  # legitimate paying user — keep Premium
        if is_admin_email(u.get("email")):
            # Owner/admin: never revoke. Their access comes from the
            # whitelist anyway, but we also leave the DB flag untouched
            # so analytics stay consistent.
            kept_admins.append(u.get("email"))
            continue
        await db.users.update_one(
            {"id": u["id"]},
            {"$set": {
                "subscription_active": False,
                "subscription_end": None,
                "premium_revoked_at": datetime.now(timezone.utc).isoformat(),
                "premium_revoked_reason": "test_period_ended",
            }}
        )
        affected.append({
            "email": u.get("email"),
            "previous_end": u.get("subscription_end"),
        })

    return {
        "revoked": len(affected),
        "users": affected,
        "kept_paying": len(paid_set),
        "kept_admins": kept_admins,
    }


@api_router.get("/admin/users-overview")
async def admin_users_overview(request: Request):
    """Quick overview of user base by plan (admin-only)."""
    _verify_admin(request.headers.get("X-Admin-Secret"))
    total = await db.users.count_documents({})
    premium = await db.users.count_documents({"subscription_active": True})
    free = total - premium
    paid = await db.payment_transactions.count_documents({"payment_status": "paid"})
    return {
        "total_users": total,
        "premium_users": premium,
        "free_users": free,
        "paid_transactions": paid,
    }

# ============== I CHING CONSULTATION ROUTES ==============
@api_router.post("/consultations", response_model=ConsultationResponse)
@limiter.limit("30/hour")  # max 30 stese/ora per IP → blocca abuso AI (€$$ Gemini)
async def create_consultation(request: Request, data: ConsultationCreate, user: dict = Depends(get_current_user)):
    # Check consultation limits
    limit_check = await check_consultation_limit(db, user)
    if not limit_check["allowed"]:
        raise HTTPException(status_code=403, detail=limit_check["message"])
    
    # Check if user can use this consultation type
    consultation_type = data.consultation_type if hasattr(data, 'consultation_type') else "deep"
    if not can_use_consultation_type(user, consultation_type):
        raise HTTPException(
            status_code=403, 
            detail="La Stesa Profonda è disponibile solo per utenti Premium. Passa a Premium o scegli la Stesa Diretta."
        )
    
    # Calculate hexagram
    hex_data = calculate_hexagram(data.coin_tosses)
    
    primary = HEXAGRAMS.get(hex_data["primary_hexagram"], {})
    derived = HEXAGRAMS.get(hex_data["derived_hexagram"], {}) if hex_data["derived_hexagram"] else None
    
    lang = user.get("language", "it")
    name_key = "name_it" if lang == "it" else "name_en"
    
    # Get traditional data
    primary_traditional = get_hexagram_traditional_data(hex_data["primary_hexagram"], lang)
    derived_traditional = get_hexagram_traditional_data(hex_data["derived_hexagram"], lang) if hex_data["derived_hexagram"] else None
    
    # Build traditional data response
    def build_traditional_response(trad_data, moving_lines, hex_num):
        trigram_above_info = get_trigram_info(trad_data.get("trigram_above", "☰"), lang)
        trigram_below_info = get_trigram_info(trad_data.get("trigram_below", "☷"), lang)
        # Use get_all_lines_text to get ALL 6 lines with is_active flag
        all_lines_texts = get_all_lines_text(hex_num, moving_lines, lang)
        
        return TraditionalData(
            sentence=trad_data.get("sentence", ""),
            image=trad_data.get("image", ""),
            commentary=trad_data.get("commentary", ""),
            trigram_above=TrigramInfo(**trigram_above_info),
            trigram_below=TrigramInfo(**trigram_below_info),
            moving_lines_text=[MovingLineText(**m) for m in all_lines_texts]
        )
    
    primary_trad_response = build_traditional_response(primary_traditional, hex_data["moving_lines"], hex_data["primary_hexagram"])
    derived_trad_response = build_traditional_response(derived_traditional, [], hex_data["derived_hexagram"]) if derived_traditional else None
    
    # Handle conversation continuation
    conversation_history = []
    conversation_depth = 0
    parent_consultation_id = data.parent_consultation_id if hasattr(data, 'parent_consultation_id') else None
    
    if parent_consultation_id:
        # Fetch the conversation history (up to last 5 consultations)
        conversation_history = await get_conversation_history(parent_consultation_id, user["id"], max_depth=5)
        conversation_depth = len(conversation_history)
    
    # Generate interpretation based on consultation type and conversation context
    consultation_type = data.consultation_type if hasattr(data, 'consultation_type') else "deep"
    topic = data.topic if hasattr(data, 'topic') else None
    interpretation = await generate_interpretation(
        hex_data, data.question, lang, consultation_type, 
        conversation_history=conversation_history,
        topic=topic
    )
    
    # Create consultation record
    consultation_id = str(uuid.uuid4())
    consultation_doc = {
        "id": consultation_id,
        "user_id": user["id"],
        "question": data.question,
        "coin_tosses": data.coin_tosses.model_dump(),
        "consultation_type": consultation_type,
        "topic": topic,
        "parent_consultation_id": parent_consultation_id,
        "conversation_depth": conversation_depth,
        "hexagram_number": hex_data["primary_hexagram"],
        "hexagram_name": primary.get(name_key, primary.get("name", "")),
        "hexagram_chinese": primary.get("name", ""),
        "hexagram_symbol": get_hexagram_symbol(hex_data["lines"]),
        "derived_hexagram_number": hex_data["derived_hexagram"],
        "derived_hexagram_name": derived.get(name_key, derived.get("name", "")) if derived else None,
        "derived_hexagram_chinese": derived.get("name", "") if derived else None,
        "moving_lines": hex_data["moving_lines"],
        "traditional_data": primary_trad_response.model_dump() if primary_trad_response else None,
        "derived_traditional_data": derived_trad_response.model_dump() if derived_trad_response else None,
        "interpretation": interpretation,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.consultations.insert_one(consultation_doc)

    # Se l'utente sta usando il gettone-prova, scala 1 credito (idempotente
    # rispetto al record consultation_doc: una consultazione → un credito).
    await consume_trial_credit_if_applicable(db, user)

    # Check and award badges
    new_badges = await check_and_award_badges(db, user["id"], consultation_doc)

    response = ConsultationResponse(
        id=consultation_id,
        question=data.question,
        hexagram_number=hex_data["primary_hexagram"],
        hexagram_name=primary.get(name_key, primary.get("name", "")),
        hexagram_chinese=primary.get("name", ""),
        hexagram_symbol=get_hexagram_symbol(hex_data["lines"]),
        derived_hexagram_number=hex_data["derived_hexagram"],
        derived_hexagram_name=derived.get(name_key, derived.get("name", "")) if derived else None,
        derived_hexagram_chinese=derived.get("name", "") if derived else None,
        moving_lines=hex_data["moving_lines"],
        traditional_data=primary_trad_response,
        derived_traditional_data=derived_trad_response,
        interpretation=interpretation,
        created_at=consultation_doc["created_at"],
        consultation_type=consultation_type,
        parent_consultation_id=parent_consultation_id,
        conversation_depth=conversation_depth
    )
    
    # Log new badges for potential frontend notification
    if new_badges:
        logger.info(f"User {user['id']} earned badges: {[b['id'] for b in new_badges]}")
    
    return response

@api_router.get("/consultations", response_model=List[ConsultationResponse])
async def get_consultations(user: dict = Depends(get_current_user)):
    consultations = await db.consultations.find(
        {"user_id": user["id"]}, 
        {"_id": 0}
    ).sort("created_at", -1).to_list(100)
    
    lang = user.get("language", "it")
    result = []
    for c in consultations:
        # Enrich old consultations with missing data
        c = enrich_consultation_data(c, lang)
        result.append(ConsultationResponse(**c))
    return result

@api_router.get("/consultations/{consultation_id}", response_model=ConsultationResponse)
async def get_consultation(consultation_id: str, user: dict = Depends(get_current_user)):
    consultation = await db.consultations.find_one(
        {"id": consultation_id, "user_id": user["id"]},
        {"_id": 0}
    )
    if not consultation:
        raise HTTPException(status_code=404, detail="Consultazione non trovata")
    
    lang = user.get("language", "it")
    consultation = enrich_consultation_data(consultation, lang)
    return ConsultationResponse(**consultation)

@api_router.delete("/consultations/{consultation_id}")
async def delete_consultation(consultation_id: str, user: dict = Depends(get_current_user)):
    """Delete a consultation"""
    result = await db.consultations.delete_one(
        {"id": consultation_id, "user_id": user["id"]}
    )
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Consultazione non trovata")
    return {"message": "Consultazione eliminata"}

@api_router.post("/consultations/synthesis", response_model=ConsultationResponse)
async def create_synthesis_consultation(data: SynthesisRequest, user: dict = Depends(get_current_user)):
    """
    Create a synthesis consultation from multiple existing consultations.
    This generates a new interpretation that combines and analyzes the selected readings.
    """
    if len(data.consultation_ids) < 2:
        raise HTTPException(status_code=400, detail="Seleziona almeno 2 consultazioni")
    
    if len(data.consultation_ids) > 5:
        raise HTTPException(status_code=400, detail="Massimo 5 consultazioni per sintesi")
    
    # Fetch all selected consultations
    consultations = []
    for cid in data.consultation_ids:
        c = await db.consultations.find_one(
            {"id": cid, "user_id": user["id"]},
            {"_id": 0}
        )
        if not c:
            raise HTTPException(status_code=404, detail=f"Consultazione {cid} non trovata")
        consultations.append(c)
    
    # Sort by creation date
    consultations.sort(key=lambda x: x.get("created_at", ""))
    
    lang = user.get("language", "it")
    
    # Build synthesis prompt
    synthesis_type_labels = {
        "confirmation": "conferma o smentita" if lang == "it" else "confirmation or denial",
        "deepening": "approfondimento" if lang == "it" else "deepening",
        "clarification": "chiarimento" if lang == "it" else "clarification"
    }
    synthesis_label = synthesis_type_labels.get(data.synthesis_type, synthesis_type_labels["deepening"])
    
    # Generate synthesis interpretation
    synthesis_interpretation = await generate_synthesis_interpretation(
        consultations, 
        data.synthesis_type, 
        lang
    )
    
    # Create combined question
    questions = [c.get("question", "") for c in consultations]
    combined_question = f"[SINTESI - {synthesis_label.upper()}]\n" + "\n→ ".join(questions)
    
    # Use the most recent hexagram as the "primary" for the synthesis
    latest = consultations[-1]
    
    # Create synthesis consultation record
    consultation_id = str(uuid.uuid4())
    consultation_doc = {
        "id": consultation_id,
        "user_id": user["id"],
        "question": combined_question,
        "hexagram_number": latest.get("hexagram_number", 1),
        "hexagram_name": latest.get("hexagram_name", ""),
        "hexagram_chinese": latest.get("hexagram_chinese", ""),
        "hexagram_symbol": latest.get("hexagram_symbol", ""),
        "derived_hexagram_number": latest.get("derived_hexagram_number"),
        "derived_hexagram_name": latest.get("derived_hexagram_name"),
        "derived_hexagram_chinese": latest.get("derived_hexagram_chinese"),
        "moving_lines": latest.get("moving_lines", []),
        "traditional_data": latest.get("traditional_data"),
        "derived_traditional_data": latest.get("derived_traditional_data"),
        "interpretation": synthesis_interpretation,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "is_synthesis": True,
        "linked_consultation_ids": data.consultation_ids,
        "synthesis_type": data.synthesis_type
    }
    
    await db.consultations.insert_one(consultation_doc)
    
    return ConsultationResponse(
        id=consultation_id,
        question=combined_question,
        hexagram_number=latest.get("hexagram_number", 1),
        hexagram_name=latest.get("hexagram_name", ""),
        hexagram_chinese=latest.get("hexagram_chinese", ""),
        hexagram_symbol=latest.get("hexagram_symbol", ""),
        derived_hexagram_number=latest.get("derived_hexagram_number"),
        derived_hexagram_name=latest.get("derived_hexagram_name"),
        derived_hexagram_chinese=latest.get("derived_hexagram_chinese"),
        moving_lines=latest.get("moving_lines", []),
        traditional_data=TraditionalData(**latest["traditional_data"]) if latest.get("traditional_data") else None,
        derived_traditional_data=TraditionalData(**latest["derived_traditional_data"]) if latest.get("derived_traditional_data") else None,
        interpretation=synthesis_interpretation,
        created_at=consultation_doc["created_at"],
        is_synthesis=True,
        linked_consultation_ids=data.consultation_ids,
        synthesis_type=data.synthesis_type
    )

async def generate_synthesis_interpretation(consultations: List[dict], synthesis_type: str, language: str) -> str:
    """Generate AI interpretation that synthesizes multiple consultations"""
    if not GEMINI_API_KEY:
        return "Interpretazione di sintesi non disponibile."

    synthesis_system = """Sei un maestro di I Ching con profonda saggezza taoista.
Il tuo compito è analizzare MULTIPLE consultazioni fatte dallo stesso consultante e creare una SINTESI che:
- Trova il filo conduttore tra le diverse stese
- Identifica conferme, contraddizioni o approfondimenti
- Offre una visione d'insieme illuminante
- Mantiene un tono rispettoso, profondo ma accessibile

Non usare elenchi puntati. Scrivi in modo fluido e narrativo.
Parla sempre in seconda persona al consultante."""

    try:
        
        # Build the consultation summaries
        summaries = []
        for i, c in enumerate(consultations, 1):
            hex_name = c.get("hexagram_name", "")
            hex_num = c.get("hexagram_number", 0)
            question = c.get("question", "")
            interpretation = c.get("interpretation", "")[:500]  # Limit length
            derived = c.get("derived_hexagram_name", "")
            moving = c.get("moving_lines", [])
            
            summary = f"""
STESA {i}:
- Domanda: {question}
- Esagramma: {hex_num}. {hex_name}
- Linee mutevoli: {moving if moving else 'Nessuna'}
- Esagramma derivato: {derived if derived else 'Nessuno'}
- Interpretazione originale (estratto): {interpretation}...
"""
            summaries.append(summary)
        
        synthesis_instructions = {
            "confirmation": {
                "it": "Analizza se le stese successive CONFERMANO o SMENTISCONO il messaggio della prima. Cerca coerenza o contraddizioni.",
                "en": "Analyze whether the subsequent readings CONFIRM or DENY the message of the first. Look for coherence or contradictions."
            },
            "deepening": {
                "it": "Approfondisci il significato complessivo, trovando connessioni nascoste tra le stese. Offri una comprensione più profonda.",
                "en": "Deepen the overall meaning, finding hidden connections between the readings. Offer a deeper understanding."
            },
            "clarification": {
                "it": "Chiarisci eventuali ambiguità, offrendo una lettura definitiva che risolva dubbi o incertezze emerse.",
                "en": "Clarify any ambiguities, offering a definitive reading that resolves doubts or uncertainties."
            }
        }
        
        instruction = synthesis_instructions.get(synthesis_type, synthesis_instructions["deepening"])
        instruction_text = instruction.get(language, instruction["it"])
        
        if language == "it":
            prompt = f"""Ecco le consultazioni I Ching da sintetizzare:

{"".join(summaries)}

ISTRUZIONI: {instruction_text}

Scrivi una SINTESI DIVINATORIA (300-500 parole) che:
1. Identifica il tema comune o l'evoluzione tra le stese
2. Analizza come gli esagrammi dialogano tra loro
3. Offre una conclusione illuminante per il consultante
4. Se ci sono linee mutevoli, considera la direzione del cambiamento

Concludi con un consiglio pratico basato sulla sintesi."""
        else:
            prompt = f"""Here are the I Ching consultations to synthesize:

{"".join(summaries)}

INSTRUCTIONS: {instruction_text}

Write a DIVINATORY SYNTHESIS (300-500 words) that:
1. Identifies the common theme or evolution between readings
2. Analyzes how the hexagrams dialogue with each other
3. Offers an illuminating conclusion for the querent
4. If there are moving lines, consider the direction of change

Conclude with practical advice based on the synthesis."""

        text = await _gemini_generate(
            system_instruction=synthesis_system,
            prompt=prompt,
            generation_config=GEMINI_DEEP_CONFIG,
        )
        return text or ("La sintesi delle tue consultazioni rivela un percorso di crescita. Gli esagrammi che hai ricevuto dialogano tra loro, suggerendo un'evoluzione del tuo cammino. Medita su come i messaggi si collegano nella tua situazione attuale." if language == "it" else "The synthesis of your consultations reveals a path of growth. The hexagrams you received dialogue with each other, suggesting an evolution of your journey. Meditate on how the messages connect in your current situation.")

    except Exception as e:
        logger.error(f"Error generating synthesis: {e}")
        if language == "it":
            return "La sintesi delle tue consultazioni rivela un percorso di crescita. Gli esagrammi che hai ricevuto dialogano tra loro, suggerendo un'evoluzione del tuo cammino. Medita su come i messaggi si collegano nella tua situazione attuale."
        return "The synthesis of your consultations reveals a path of growth. The hexagrams you received dialogue with each other, suggesting an evolution of your journey. Meditate on how the messages connect in your current situation."

def enrich_consultation_data(consultation: dict, language: str) -> dict:
    """Add missing traditional data to old consultations"""
    hex_num = consultation.get("hexagram_number")
    derived_num = consultation.get("derived_hexagram_number")
    moving_lines = consultation.get("moving_lines", [])
    
    # Add hexagram_chinese if missing
    if not consultation.get("hexagram_chinese"):
        primary = HEXAGRAMS.get(hex_num, {})
        consultation["hexagram_chinese"] = primary.get("name", "")
    
    # Add derived_hexagram_chinese if missing
    if derived_num and not consultation.get("derived_hexagram_chinese"):
        derived = HEXAGRAMS.get(derived_num, {})
        consultation["derived_hexagram_chinese"] = derived.get("name", "")
    
    # Add traditional_data if missing
    if not consultation.get("traditional_data"):
        trad_data = get_hexagram_traditional_data(hex_num, language)
        trigram_above_info = get_trigram_info(trad_data.get("trigram_above", "☰"), language)
        trigram_below_info = get_trigram_info(trad_data.get("trigram_below", "☷"), language)
        # Use get_all_lines_text to get ALL 6 lines
        all_lines_texts = get_all_lines_text(hex_num, moving_lines, language)
        
        consultation["traditional_data"] = {
            "sentence": trad_data.get("sentence", ""),
            "image": trad_data.get("image", ""),
            "commentary": trad_data.get("commentary", ""),
            "trigram_above": trigram_above_info,
            "trigram_below": trigram_below_info,
            "moving_lines_text": all_lines_texts
        }
    
    # Add derived_traditional_data if missing
    if derived_num and not consultation.get("derived_traditional_data"):
        derived_trad = get_hexagram_traditional_data(derived_num, language)
        d_trigram_above = get_trigram_info(derived_trad.get("trigram_above", "☰"), language)
        d_trigram_below = get_trigram_info(derived_trad.get("trigram_below", "☷"), language)
        
        consultation["derived_traditional_data"] = {
            "sentence": derived_trad.get("sentence", ""),
            "image": derived_trad.get("image", ""),
            "commentary": derived_trad.get("commentary", ""),
            "trigram_above": d_trigram_above,
            "trigram_below": d_trigram_below,
            "moving_lines_text": []
        }
    
    # Add synthesis fields if missing (for backward compatibility)
    if "is_synthesis" not in consultation:
        consultation["is_synthesis"] = False
    if "linked_consultation_ids" not in consultation:
        consultation["linked_consultation_ids"] = []
    if "synthesis_type" not in consultation:
        consultation["synthesis_type"] = None
    
    # Add consultation_type if missing (for backward compatibility)
    if "consultation_type" not in consultation:
        consultation["consultation_type"] = "deep"  # Default to deep for old consultations
    
    # Add conversation fields if missing
    if "parent_consultation_id" not in consultation:
        consultation["parent_consultation_id"] = None
    if "conversation_depth" not in consultation:
        consultation["conversation_depth"] = 0
    
    return consultation

# ============== SHARE CONSULTATION ==============
@api_router.post("/consultations/{consultation_id}/share")
async def create_share_link(consultation_id: str, user: dict = Depends(get_current_user)):
    """Generate a public share token for a consultation"""
    consultation = await db.consultations.find_one(
        {"id": consultation_id, "user_id": user["id"]},
        {"_id": 0}
    )
    if not consultation:
        raise HTTPException(status_code=404, detail="Consultazione non trovata")
    
    # Generate share token if not exists
    share_token = consultation.get("share_token")
    if not share_token:
        # Token di condivisione consultazione: 32 caratteri url-safe = 192 bit
        # di entropia (vs 48 bit del vecchio uuid[:12]). I bruteforce su
        # questo token diventano computazionalmente impossibili.
        share_token = secrets.token_urlsafe(24)
        await db.consultations.update_one(
            {"id": consultation_id},
            {"$set": {"share_token": share_token, "is_public": True}}
        )
    
    return {"share_token": share_token, "consultation_id": consultation_id}

@api_router.get("/shared/{share_token}")
async def get_shared_consultation(share_token: str):
    """Get a publicly shared consultation (no auth required)"""
    consultation = await db.consultations.find_one(
        {"share_token": share_token, "is_public": True},
        {"_id": 0, "user_id": 0}  # Exclude user info for privacy
    )
    if not consultation:
        raise HTTPException(status_code=404, detail="Consultazione non trovata o non condivisa")
    
    return {
        "id": consultation["id"],
        "question": consultation["question"],
        "hexagram_number": consultation["hexagram_number"],
        "hexagram_name": consultation["hexagram_name"],
        "derived_hexagram_number": consultation.get("derived_hexagram_number"),
        "derived_hexagram_name": consultation.get("derived_hexagram_name"),
        "moving_lines": consultation["moving_lines"],
        "interpretation": consultation["interpretation"],
        "created_at": consultation["created_at"]
    }

# ============== STRIPE PAYMENT ROUTES ==============
# Prices come from subscription_manager.SUBSCRIPTION_PRICES so they stay
# in one place (and match what /subscription/status exposes to the UI).
SUBSCRIPTION_PRICE = SUBSCRIPTION_PRICES["monthly"]["price"]  # legacy export


@api_router.post("/payments/checkout")
async def create_checkout(data: CheckoutRequest, request: Request, user: dict = Depends(get_current_user)):
    if not STRIPE_API_KEY:
        raise HTTPException(
            status_code=503,
            detail="Pagamenti non ancora configurati. Riprova più tardi."
        )

    # Map legacy aliases to current SKUs
    plan_type = (data.plan_type or "base_monthly").lower()
    legacy_map = {"monthly": "base_monthly", "yearly": "base_yearly"}
    plan_type = legacy_map.get(plan_type, plan_type)
    if plan_type not in SUBSCRIPTION_PRICES:
        raise HTTPException(
            status_code=400,
            detail="plan_type non valido. Valori ammessi: trial_pack | base_monthly | base_yearly | fitness_monthly | fitness_yearly"
        )
    plan_cfg = SUBSCRIPTION_PRICES[plan_type]
    price_eur = float(plan_cfg["price"])
    is_trial = plan_type == "trial_pack"
    # I trial non hanno scadenza a tempo: usano crediti. Per gli altri, default 30/365 giorni.
    duration_days = 0 if is_trial else int(plan_cfg.get("duration_days") or (365 if "yearly" in plan_type else 30))
    plan_name = plan_cfg.get("plan", "base")  # 'base' / 'fitness_coaching' / 'trial_pack'
    trial_credits = int(plan_cfg.get("trial_credits") or 0) if is_trial else 0
    product_label = plan_cfg.get(
        "label_it",
        "I Ching del Benessere — Abbonamento"
    )

    stripe_lib.api_key = STRIPE_API_KEY

    # SICUREZZA: `origin_url` arriva dal client, quindi non ci fidiamo. Un
    # attaccante puo' sostituirlo con "https://phishing.evil" e reindirizzare
    # l'utente dopo il pagamento verso un sito controllato. Whitelist:
    # accettiamo solo origini che matchano APP_URL o localhost dev.
    from urllib.parse import urlparse
    _allowed_hosts = {
        urlparse(os.environ.get("APP_URL", "https://www.chingbenessere.it")).netloc,
        "www.chingbenessere.it",
        "chingbenessere.it",
        "localhost:3000",
        "localhost:3001",
    }
    parsed = urlparse(data.origin_url or "")
    if parsed.scheme not in ("http", "https") or parsed.netloc not in _allowed_hosts:
        # Fallback silente all'origine ufficiale: non blocchiamo l'acquisto,
        # ma redirigiamo alla home canonica dopo il pagamento.
        base = os.environ.get("APP_URL", "https://www.chingbenessere.it").rstrip("/")
    else:
        base = f"{parsed.scheme}://{parsed.netloc}"
    success_url = f"{base}/payment/success?session_id={{CHECKOUT_SESSION_ID}}"
    cancel_url = f"{base}/subscription"

    session = stripe_lib.checkout.Session.create(
        payment_method_types=["card"],
        line_items=[{
            "price_data": {
                "currency": plan_cfg["currency"].lower(),
                "product_data": {"name": product_label},
                "unit_amount": int(round(price_eur * 100)),
            },
            "quantity": 1,
        }],
        mode="payment",
        success_url=success_url,
        cancel_url=cancel_url,
        metadata={
            "user_id": user["id"],
            "user_email": user["email"],
            "type": f"{plan_type}_purchase" if is_trial else f"{plan_type}_subscription",
            "plan": plan_name,  # 'base' / 'fitness_coaching' / 'trial_pack'
            "duration_days": str(duration_days),
            "trial_credits": str(trial_credits),
        }
    )

    transaction_doc = {
        "id": str(uuid.uuid4()),
        "session_id": session.id,
        "user_id": user["id"],
        "amount": price_eur,
        "currency": plan_cfg["currency"].lower(),
        "plan_type": plan_type,
        "plan": plan_name,
        "duration_days": duration_days,
        "trial_credits": trial_credits,
        "is_trial": is_trial,
        "status": "pending",
        "payment_status": "initiated",
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.payment_transactions.insert_one(transaction_doc)

    return {"url": session.url, "session_id": session.id}

@api_router.post("/subscription/cancel")
async def cancel_subscription(user: dict = Depends(get_current_user)):
    """
    User-initiated subscription cancellation.

    We DON'T immediately revoke access — the user has already paid for the
    current period, so Premium stays active until subscription_end. We just
    set `auto_renew=False` and `cancellation_requested_at` so:
      - the UI can show "Annullato, attivo fino a DD/MM/YYYY"
      - no future renewal is triggered
      - we have an audit trail

    Once subscription_end is in the past, get_user_plan() naturally returns 'free'.
    """
    # Admins / whitelist users have lifetime free Premium — cancel makes no sense
    if is_admin_email(user.get("email")):
        raise HTTPException(
            status_code=400,
            detail="Gli account amministratori hanno accesso permanente — non c'è nulla da disdire."
        )

    if not user.get("subscription_active"):
        raise HTTPException(status_code=400, detail="Non hai un abbonamento attivo da disdire.")

    now_iso = datetime.now(timezone.utc).isoformat()
    await db.users.update_one(
        {"id": user["id"]},
        {"$set": {
            "auto_renew": False,
            "cancellation_requested_at": now_iso,
        }}
    )
    sub_end = user.get("subscription_end")

    # Email di conferma disdetta — best effort
    try:
        mailer.send_subscription_cancelled(
            to=user.get("email") or "",
            user_name=user.get("name") or "",
            active_until=sub_end,
        )
    except Exception as e:  # noqa: BLE001
        logger.warning("Cancellation email failed for %s: %s", user.get("email"), e)
    return {
        "message": "Abbonamento disdetto. Continuerai ad avere accesso Premium fino alla scadenza.",
        "active_until": sub_end,
        "cancelled_at": now_iso,
    }


@api_router.post("/subscription/withdraw")
async def withdraw_subscription(user: dict = Depends(get_current_user)):
    """
    Right of withdrawal (art. 52 Codice del Consumo / 14 days from purchase).

    Triggers immediate Premium revocation. Refund is NOT issued automatically —
    we just flag the request and the admin processes it manually via Stripe
    (this lets us subtract any consultations already consumed, as the law allows).
    """
    if is_admin_email(user.get("email")):
        raise HTTPException(status_code=400, detail="Gli admin non possono esercitare il recesso.")

    # Find the latest paid transaction
    last_paid = await db.payment_transactions.find_one(
        {"user_id": user["id"], "payment_status": "paid"},
        {"_id": 0},
        sort=[("created_at", -1)],
    )
    if not last_paid:
        raise HTTPException(status_code=400, detail="Nessun pagamento da rimborsare.")

    # 14-day window check. Se la data acquisto non e' parsabile RIFIUTIAMO
    # invece di assumere "oggi" (vecchio bug: fallback a datetime.now()
    # rendeva il recesso SEMPRE valido anche per acquisti vecchi).
    try:
        purchase_dt = datetime.fromisoformat(str(last_paid.get("created_at", "")).replace("Z", "+00:00"))
    except Exception:
        raise HTTPException(
            status_code=400,
            detail="Data di acquisto non leggibile. Contatta l'assistenza per il recesso.",
        )
    elapsed_days = (datetime.now(timezone.utc) - purchase_dt).days
    if elapsed_days > 14:
        raise HTTPException(
            status_code=400,
            detail=f"Il diritto di recesso è scaduto (acquisto avvenuto {elapsed_days} giorni fa, limite 14)."
        )

    now_iso = datetime.now(timezone.utc).isoformat()

    # Revoke access immediately (refund handled manually by admin)
    await db.users.update_one(
        {"id": user["id"]},
        {"$set": {
            "subscription_active": False,
            "subscription_end": None,
            "withdrawal_requested_at": now_iso,
            "withdrawal_session_id": last_paid["session_id"],
            "auto_renew": False,
        }}
    )
    # Record the request so admin can process the refund
    await db.refund_requests.insert_one({
        "id": str(uuid.uuid4()),
        "user_id": user["id"],
        "email": user["email"],
        "session_id": last_paid["session_id"],
        "amount": last_paid.get("amount"),
        "currency": last_paid.get("currency"),
        "requested_at": now_iso,
        "status": "pending",
        "purchase_date": last_paid["created_at"],
    })

    # Email conferma recesso — best effort
    try:
        mailer.send_withdrawal_confirmed(
            to=user.get("email") or "",
            user_name=user.get("name") or "",
        )
    except Exception as e:  # noqa: BLE001
        logger.warning("Withdrawal email failed for %s: %s", user.get("email"), e)

    return {
        "message": "Recesso esercitato. Il rimborso verrà processato entro 14 giorni lavorativi.",
        "withdrawn_at": now_iso,
    }


@api_router.get("/payments/status/{session_id}")
async def get_payment_status(session_id: str, user: dict = Depends(get_current_user)):
    stripe_lib.api_key = STRIPE_API_KEY

    session = stripe_lib.checkout.Session.retrieve(session_id)

    # Update transaction record
    if session.payment_status == "paid":
        existing = await db.payment_transactions.find_one({
            "session_id": session_id,
            "payment_status": "paid"
        })

        if not existing:
            await db.payment_transactions.update_one(
                {"session_id": session_id},
                {"$set": {
                    "status": session.status,
                    "payment_status": session.payment_status,
                    "updated_at": datetime.now(timezone.utc).isoformat()
                }}
            )

            # Activate subscription for the chosen plan duration
            txn = await db.payment_transactions.find_one({"session_id": session_id}, {"_id": 0})
            days = (txn or {}).get("duration_days") or 30
            plan_name = (txn or {}).get("plan") or "base"
            subscription_end = datetime.now(timezone.utc) + timedelta(days=int(days))
            await db.users.update_one(
                {"id": user["id"]},
                {"$set": {
                    "subscription_active": True,
                    "subscription_end": subscription_end.isoformat(),
                    "subscription_plan": plan_name,
                }}
            )

    return {
        "status": session.status,
        "payment_status": session.payment_status,
        "amount_total": session.amount_total,
        "currency": session.currency
    }

@api_router.post("/webhook/stripe")
async def stripe_webhook(request: Request):
    body = await request.body()
    sig_header = request.headers.get("Stripe-Signature", "")
    
    # SICUREZZA: il webhook DEVE verificare la firma Stripe. Senza
    # verifica un attaccante puo' fingere `checkout.session.completed`
    # con metadata.user_id arbitrario e ottenere Premium gratis per
    # qualsiasi utente.
    #
    # Vecchio codice: if STRIPE_WEBHOOK_SECRET and sig_header: verify
    #                 else: accept-without-verify  ← bypass se config rotta
    #
    # Nuovo: rifiutiamo TUTTO se non possiamo verificare. L'unica
    # eccezione è il flag di sviluppo STRIPE_WEBHOOK_INSECURE=1, usato
    # solo per test locali col Stripe CLI listen.
    is_dev_insecure = os.environ.get('STRIPE_WEBHOOK_INSECURE') == '1'
    if not STRIPE_WEBHOOK_SECRET and not is_dev_insecure:
        logger.error("Webhook Stripe ricevuto ma STRIPE_WEBHOOK_SECRET non configurato — rifiuto.")
        raise HTTPException(status_code=503, detail="Webhook non configurato.")
    if not sig_header and not is_dev_insecure:
        raise HTTPException(status_code=400, detail="Stripe-Signature mancante.")

    try:
        stripe_lib.api_key = STRIPE_API_KEY
        if STRIPE_WEBHOOK_SECRET and sig_header:
            try:
                event = stripe_lib.Webhook.construct_event(body, sig_header, STRIPE_WEBHOOK_SECRET)
            except stripe_lib.error.SignatureVerificationError:
                logger.warning("Webhook Stripe con firma INVALIDA — rifiuto.")
                raise HTTPException(status_code=400, detail="Firma webhook non valida.")
        elif is_dev_insecure:
            # Solo in dev: bypassa verifica firma per testing locale
            logger.warning("⚠️  STRIPE_WEBHOOK_INSECURE=1 attivo — la firma NON verra' verificata. Mai usare in produzione!")
            event = stripe_lib.Event.construct_from(
                stripe_lib.util.convert_to_stripe_object(
                    stripe_lib.util.json.loads(body)
                ),
                stripe_lib.api_key
            )
        else:
            raise HTTPException(status_code=400, detail="Webhook non verificabile.")

        if event["type"] == "checkout.session.completed":
            session = event["data"]["object"]
            if session.get("payment_status") == "paid":
                user_id = session.get("metadata", {}).get("user_id")
                if user_id:
                    existing = await db.payment_transactions.find_one({
                        "session_id": session["id"],
                        "payment_status": "paid"
                    })
                    if not existing:
                        await db.payment_transactions.update_one(
                            {"session_id": session["id"]},
                            {"$set": {
                                "status": "complete",
                                "payment_status": "paid",
                                "updated_at": datetime.now(timezone.utc).isoformat()
                            }}
                        )
                        # Use duration + plan from transaction record
                        txn = await db.payment_transactions.find_one(
                            {"session_id": session["id"]}, {"_id": 0}
                        )
                        plan_name = (txn or {}).get("plan") or "base"
                        plan_type = (txn or {}).get("plan_type") or ""
                        amount = float((txn or {}).get("amount") or 0)
                        is_trial = plan_name == "trial_pack" or bool((txn or {}).get("is_trial"))
                        days_credited = 0
                        trial_credits_credited = 0
                        if is_trial:
                            # GETTONE PROVA: accredita N consultazioni, niente
                            # subscription_active. Funziona a crediti.
                            trial_credits_credited = int((txn or {}).get("trial_credits") or 3)
                            await grant_trial_pack(db, user_id, credits=trial_credits_credited)
                        else:
                            days_credited = int((txn or {}).get("duration_days") or 30)
                            subscription_end = datetime.now(timezone.utc) + timedelta(days=days_credited)
                            await db.users.update_one(
                                {"id": user_id},
                                {"$set": {
                                    "subscription_active": True,
                                    "subscription_end": subscription_end.isoformat(),
                                    "subscription_plan": plan_name,
                                }}
                            )

                        # Email ricevuta — best effort, non blocca. Etichetta
                        # leggibile per Stripe Checkout + email l'abbiamo da txn metadata
                        # oppure dal record utente.
                        try:
                            target_user = await db.users.find_one({"id": user_id}, {"_id": 0, "email": 1, "name": 1})
                            if target_user and target_user.get("email"):
                                plan_label_map = {
                                    "trial_pack": "Gettone Prova — 3 consultazioni",
                                    "base_monthly": "Base — Mensile",
                                    "base_yearly": "Base — Annuale",
                                    "fitness_monthly": "Benessere Fisico — Mensile",
                                    "fitness_yearly": "Benessere Fisico — Annuale",
                                }
                                plan_label = plan_label_map.get(plan_type, plan_type or "Abbonamento")
                                mailer.send_payment_receipt(
                                    to=target_user["email"],
                                    user_name=target_user.get("name") or "",
                                    plan_label=plan_label,
                                    amount_eur=amount,
                                    is_trial=is_trial,
                                    duration_days=days_credited,
                                    trial_credits=trial_credits_credited,
                                )
                        except Exception as e:  # noqa: BLE001
                            logger.warning("Receipt email failed for user %s: %s", user_id, e)

        return {"received": True}
    except HTTPException:
        # Propaga i 4xx/5xx generati esplicitamente (firma invalida, secret
        # mancante). Prima del fix un `except Exception` piatto ingoiava
        # anche gli HTTPException e ritornava 200 - un attaccante poteva
        # inviare webhook fasulli e ottenere Premium gratis.
        raise
    except Exception as e:
        logger.error(f"Webhook error: {e}")
        # Errori "interni" (parse, DB down) -> 200 per non far ritentare
        # Stripe in loop. Se necessario riprocessiamo manualmente da
        # Stripe dashboard.
        return {"received": True}

# ============== HEXAGRAM INFO ROUTE ==============
@api_router.get("/hexagrams")
async def get_hexagrams():
    return HEXAGRAMS

@api_router.get("/hexagrams/{number}")
async def get_hexagram(number: int):
    if number < 1 or number > 64:
        raise HTTPException(status_code=404, detail="Esagramma non trovato")
    return HEXAGRAMS.get(number, {})


# ============== SUBSCRIPTION & LIMITS ==============

@api_router.get("/subscription/status")
async def get_subscription_status(request: Request):
    """Get user's subscription status and limits"""
    user = await get_current_user(request)
    plan = get_user_plan(user)
    limits = get_plan_limits(plan)
    
    # Get consultation count this month.
    # created_at e' salvato come stringa ISO su insert (server.py:2235), quindi
    # il confronto con datetime non matcha per BSON type ordering: usare
    # sempre .isoformat() per allineare i tipi.
    start_of_month = datetime.now(timezone.utc).replace(day=1, hour=0, minute=0, second=0, microsecond=0).isoformat()
    monthly_count = await db.consultations.count_documents({
        "user_id": user["id"],
        "created_at": {"$gte": start_of_month}
    })
    
    remaining = limits["monthly_consultations"] - monthly_count if limits["monthly_consultations"] != -1 else -1
    
    # Determine if user is still within the 14-day withdrawal window
    within_withdrawal_window = False
    last_paid = await db.payment_transactions.find_one(
        {"user_id": user["id"], "payment_status": "paid"},
        {"_id": 0, "created_at": 1, "amount": 1},
        sort=[("created_at", -1)],
    )
    if last_paid:
        try:
            purchase_dt = datetime.fromisoformat(last_paid["created_at"].replace("Z", "+00:00"))
            within_withdrawal_window = (datetime.now(timezone.utc) - purchase_dt).days <= 14
        except Exception:
            pass

    try:
        trial_credits_remaining = int(user.get("trial_credits_remaining") or 0)
    except (TypeError, ValueError):
        trial_credits_remaining = 0
    try:
        trial_credits_consumed = int(user.get("trial_credits_consumed_total") or 0)
    except (TypeError, ValueError):
        trial_credits_consumed = 0

    return {
        "plan": plan,
        "limits": limits,
        "usage": {
            "monthly_consultations": monthly_count,
            "remaining": remaining
        },
        "trial": {
            "active": plan == "trial_pack",
            "credits_remaining": trial_credits_remaining,
            "credits_consumed": trial_credits_consumed,
            "purchased_at": user.get("trial_purchased_at"),
        },
        "subscription_end": user.get("subscription_end"),
        "auto_renew": user.get("auto_renew", True),
        "cancellation_requested_at": user.get("cancellation_requested_at"),
        "is_admin": is_admin_email(user.get("email")),
        "within_withdrawal_window": within_withdrawal_window,
        "last_purchase_at": last_paid.get("created_at") if last_paid else None,
        "last_purchase_amount": last_paid.get("amount") if last_paid else None,
        "prices": SUBSCRIPTION_PRICES
    }


@api_router.get("/subscription/check-limit")
async def check_limit(request: Request):
    """Check if user can make a consultation"""
    user = await get_current_user(request)
    result = await check_consultation_limit(db, user)
    result["plan"] = get_user_plan(user)
    return result


# ============== DAILY HEXAGRAM ==============

@api_router.get("/daily-hexagram")
async def get_daily_hexagram(request: Request):
    """Get the hexagram of the day"""
    try:
        user = await get_current_user(request)
        lang = user.get("language", "it")
    except:
        lang = "it"
    
    hex_number = get_daily_hexagram_number()
    hex_data = HEXAGRAMS.get(hex_number, {})
    extended = get_extended_hexagram_data(hex_number, lang)
    traditional = get_hexagram_traditional_data(hex_number, lang)
    lunar = get_lunar_phase()
    
    name_key = "name_it" if lang == "it" else "name_en"
    
    # Generate a short daily message
    daily_messages_it = [
        f"Oggi il Tao ti parla attraverso {hex_data.get(name_key)}. Lascia che la sua energia ti guidi.",
        f"L'esagramma {hex_number} illumina il tuo cammino oggi. Ascolta il suo messaggio.",
        f"{hex_data.get(name_key)} ti accompagna in questa giornata. Cosa vuole insegnarti?",
    ]
    daily_messages_en = [
        f"Today the Tao speaks to you through {hex_data.get(name_key)}. Let its energy guide you.",
        f"Hexagram {hex_number} illuminates your path today. Listen to its message.",
        f"{hex_data.get(name_key)} accompanies you on this day. What does it want to teach you?",
    ]
    
    import random as _rnd
    # Istanza locale (non random.seed globale) per evitare race con altre coroutine
    _daily_rng = _rnd.Random(int(datetime.now(timezone.utc).strftime("%Y%m%d")))
    message = _daily_rng.choice(daily_messages_it if lang == "it" else daily_messages_en)
    
    return {
        "date": datetime.now(timezone.utc).strftime("%Y-%m-%d"),
        "hexagram_number": hex_number,
        "hexagram_name": hex_data.get(name_key, ""),
        "hexagram_chinese": hex_data.get("name", ""),
        "sentence": extended.get("giudizio", traditional.get("sentence", "")),
        "image": extended.get("immagine", traditional.get("image", "")),
        "daily_message": message,
        "lunar_phase": lunar
    }


# ============== LUNAR CALENDAR ==============

@api_router.get("/lunar-calendar")
async def get_lunar_calendar():
    """Get current lunar phase and I Ching advice"""
    phase = get_lunar_phase()
    return phase


# ============== I CHING LIBRARY ==============

@api_router.get("/library/hexagrams")
async def get_library_hexagrams(request: Request):
    """Get all 64 hexagrams for the library"""
    try:
        user = await get_current_user(request)
        lang = user.get("language", "it")
    except:
        lang = "it"
    
    name_key = "name_it" if lang == "it" else "name_en"
    
    hexagrams_list = []
    for num in range(1, 65):
        hex_data = HEXAGRAMS.get(num, {})
        extended = ICHING_EXTENDED.get(num, {})
        
        hexagrams_list.append({
            "number": num,
            "name": hex_data.get(name_key, ""),
            "chinese": hex_data.get("name", ""),
            "trigram_above": hex_data.get("trigram_top", ""),
            "trigram_below": hex_data.get("trigram_bottom", ""),
            "giudizio": extended.get("giudizio", "")[:100] + "..." if extended.get("giudizio") else ""
        })
    
    return hexagrams_list


@api_router.get("/library/hexagrams/{number}")
async def get_library_hexagram_detail(number: int, request: Request):
    """Get detailed info for a specific hexagram"""
    if number < 1 or number > 64:
        raise HTTPException(status_code=404, detail="Esagramma non trovato")
    
    try:
        user = await get_current_user(request)
        lang = user.get("language", "it")
    except:
        lang = "it"
    
    hex_data = HEXAGRAMS.get(number, {})
    extended = ICHING_EXTENDED.get(number, {})
    traditional = get_hexagram_traditional_data(number, lang)
    
    name_key = "name_it" if lang == "it" else "name_en"
    
    # Get all 6 lines
    lines = []
    for line_num in range(1, 7):
        line_data = extended.get("linee", {}).get(line_num, {})
        lines.append({
            "position": line_num,
            "text": line_data.get("testo", ""),
            "meaning": line_data.get("significato", "")
        })
    
    # Get trigram info
    trigram_above = get_trigram_info(traditional.get("trigram_above", hex_data.get("trigram_top", "☰")), lang)
    trigram_below = get_trigram_info(traditional.get("trigram_below", hex_data.get("trigram_bottom", "☷")), lang)
    
    return {
        "number": number,
        "name": hex_data.get(name_key, ""),
        "chinese": hex_data.get("name", ""),
        "chinese_name": extended.get("nome_cinese", ""),
        "giudizio": extended.get("giudizio", traditional.get("sentence", "")),
        "immagine": extended.get("immagine", traditional.get("image", "")),
        "commento": extended.get("commento", traditional.get("commentary", "")),
        "trigram_above": trigram_above,
        "trigram_below": trigram_below,
        "lines": lines
    }


@api_router.get("/library/trigrams")
async def get_library_trigrams(request: Request):
    """Get all 8 trigrams info"""
    try:
        user = await get_current_user(request)
        lang = user.get("language", "it")
    except:
        lang = "it"
    
    trigrams_list = []
    for symbol in TRIGRAMS.keys():
        info = get_trigram_info(symbol, lang)
        trigrams_list.append(info)
    
    return trigrams_list


# ============== PERSONAL DIARY / NOTES ==============

@api_router.post("/notes")
async def create_note(data: NoteCreate, request: Request):
    """Add a personal note to a consultation (Premium only)"""
    user = await get_current_user(request)
    
    # Check if premium
    plan = get_user_plan(user)
    if plan != "premium":
        raise HTTPException(status_code=403, detail="Funzionalità Premium. Abbonati per aggiungere note personali.")
    
    # Verify consultation exists and belongs to user
    consultation = await db.consultations.find_one({
        "id": data.consultation_id,
        "user_id": user["id"]
    })
    if not consultation:
        raise HTTPException(status_code=404, detail="Consultazione non trovata")
    
    note_id = str(uuid.uuid4())
    note_doc = {
        "id": note_id,
        "user_id": user["id"],
        "consultation_id": data.consultation_id,
        "content": data.content,
        "mood": data.mood,
        "tags": data.tags or [],
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.notes.insert_one(note_doc)
    
    return {"id": note_id, "message": "Nota salvata"}


@api_router.get("/notes")
async def get_user_notes(request: Request, consultation_id: str = None):
    """Get user's notes, optionally filtered by consultation"""
    user = await get_current_user(request)
    
    query = {"user_id": user["id"]}
    if consultation_id:
        query["consultation_id"] = consultation_id
    
    notes = await db.notes.find(query, {"_id": 0}).sort("created_at", -1).to_list(100)
    return notes


@api_router.put("/notes/{note_id}")
async def update_note(note_id: str, data: NoteUpdate, request: Request):
    """Update a note"""
    user = await get_current_user(request)
    
    note = await db.notes.find_one({"id": note_id, "user_id": user["id"]})
    if not note:
        raise HTTPException(status_code=404, detail="Nota non trovata")
    
    update_data = {"updated_at": datetime.now(timezone.utc).isoformat()}
    if data.content is not None:
        update_data["content"] = data.content
    if data.mood is not None:
        update_data["mood"] = data.mood
    if data.tags is not None:
        update_data["tags"] = data.tags
    
    await db.notes.update_one({"id": note_id}, {"$set": update_data})
    return {"message": "Nota aggiornata"}


@api_router.delete("/notes/{note_id}")
async def delete_note(note_id: str, request: Request):
    """Delete a note"""
    user = await get_current_user(request)
    
    result = await db.notes.delete_one({"id": note_id, "user_id": user["id"]})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Nota non trovata")
    
    return {"message": "Nota eliminata"}


# ============== STATISTICS ==============

@api_router.get("/statistics")
async def get_user_statistics(request: Request):
    """Get user's consultation statistics (Premium only)"""
    user = await get_current_user(request)
    
    # Check if premium for full stats
    plan = get_user_plan(user)
    
    # Basic stats for everyone
    total_consultations = await db.consultations.count_documents({"user_id": user["id"]})
    
    # Get level info
    level_info = get_user_level(total_consultations)
    
    # Get badges
    user_badges = user.get("badges", [])
    badges_detail = [b for b in BADGES if b["id"] in user_badges]
    
    basic_stats = {
        "total_consultations": total_consultations,
        "level": level_info,
        "badges": badges_detail,
        "plan": plan
    }
    
    if plan != "premium":
        basic_stats["premium_required"] = True
        basic_stats["message"] = "Abbonati a Premium per vedere le statistiche complete"
        return basic_stats
    
    # Premium stats
    # Most frequent hexagrams
    pipeline = [
        {"$match": {"user_id": user["id"]}},
        {"$group": {"_id": "$hexagram_number", "count": {"$sum": 1}, "name": {"$first": "$hexagram_name"}}},
        {"$sort": {"count": -1}},
        {"$limit": 5}
    ]
    frequent_hexagrams = await db.consultations.aggregate(pipeline).to_list(5)
    
    # Topics distribution
    pipeline_topics = [
        {"$match": {"user_id": user["id"], "topic": {"$ne": None}}},
        {"$group": {"_id": "$topic", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}}
    ]
    topics_distribution = await db.consultations.aggregate(pipeline_topics).to_list(10)
    
    # Moving lines frequency
    pipeline_lines = [
        {"$match": {"user_id": user["id"]}},
        {"$unwind": "$moving_lines"},
        {"$group": {"_id": "$moving_lines", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}}
    ]
    lines_frequency = await db.consultations.aggregate(pipeline_lines).to_list(6)
    
    # Unique hexagrams encountered
    unique_hexagrams = await db.consultations.distinct("hexagram_number", {"user_id": user["id"]})
    
    # Consultations by type
    direct_count = await db.consultations.count_documents({"user_id": user["id"], "consultation_type": "direct"})
    deep_count = await db.consultations.count_documents({"user_id": user["id"], "consultation_type": "deep"})
    
    # Monthly trend (last 6 months)
    six_months_ago = datetime.now(timezone.utc) - timedelta(days=180)
    pipeline_monthly = [
        {"$match": {"user_id": user["id"], "created_at": {"$gte": six_months_ago}}},
        {"$group": {
            "_id": {"$substr": ["$created_at", 0, 7]},  # YYYY-MM
            "count": {"$sum": 1}
        }},
        {"$sort": {"_id": 1}}
    ]
    monthly_trend = await db.consultations.aggregate(pipeline_monthly).to_list(6)
    
    return {
        **basic_stats,
        "premium_required": False,
        "frequent_hexagrams": [{"number": h["_id"], "name": h["name"], "count": h["count"]} for h in frequent_hexagrams],
        "topics_distribution": {t["_id"]: t["count"] for t in topics_distribution},
        "lines_frequency": {str(l["_id"]): l["count"] for l in lines_frequency},
        "unique_hexagrams_count": len(unique_hexagrams),
        "unique_hexagrams": unique_hexagrams,
        "consultation_types": {"direct": direct_count, "deep": deep_count},
        "monthly_trend": [{"month": m["_id"], "count": m["count"]} for m in monthly_trend]
    }


# ============== GUIDED PATHS ==============

@api_router.get("/paths")
async def get_guided_paths(request: Request):
    """Get available guided paths"""
    try:
        user = await get_current_user(request)
        lang = user.get("language", "it")
    except:
        lang = "it"
    
    name_key = "name_it" if lang == "it" else "name_en"
    desc_key = "description_it" if lang == "it" else "description_en"
    
    paths_list = []
    for path_id, path in GUIDED_PATHS.items():
        paths_list.append({
            "id": path["id"],
            "name": path[name_key],
            "description": path[desc_key],
            "emoji": path["emoji"],
            "total_steps": len(path["steps"])
        })
    
    return paths_list


@api_router.get("/paths/unread-count")
async def get_unread_paths_count(request: Request):
    """Get count of unread completed paths for notification badge.
    NOTE: Must be defined BEFORE /paths/{path_id} so FastAPI doesn't
    treat 'unread-count' as a path_id parameter."""
    user = await get_current_user(request)

    count = await db.completed_paths.count_documents({
        "user_id": user["id"],
        "is_read": False
    })

    return {"count": count}


@api_router.get("/paths/{path_id}")
async def get_path_detail(path_id: str, request: Request):
    """Get details of a specific guided path"""
    user = await get_current_user(request)
    lang = user.get("language", "it")
    
    path = GUIDED_PATHS.get(path_id)
    if not path:
        raise HTTPException(status_code=404, detail="Percorso non trovato")
    
    name_key = "name_it" if lang == "it" else "name_en"
    desc_key = "description_it" if lang == "it" else "description_en"
    question_key = "question_it" if lang == "it" else "question_en"
    
    # Get user's progress on this path
    user_path = await db.user_paths.find_one({
        "user_id": user["id"],
        "path_id": path_id
    })
    
    completed_steps = user_path.get("completed_steps", []) if user_path else []
    
    steps = []
    for step in path["steps"]:
        steps.append({
            "day": step["day"],
            "question": step[question_key],
            "completed": step["day"] in completed_steps
        })
    
    return {
        "id": path["id"],
        "name": path[name_key],
        "description": path[desc_key],
        "emoji": path["emoji"],
        "steps": steps,
        "total_steps": len(path["steps"]),
        "completed_steps": len(completed_steps),
        "started": user_path is not None,
        "started_at": user_path.get("started_at") if user_path else None
    }


@api_router.post("/paths/{path_id}/start")
async def start_path(path_id: str, request: Request):
    """Start a guided path"""
    user = await get_current_user(request)
    
    path = GUIDED_PATHS.get(path_id)
    if not path:
        raise HTTPException(status_code=404, detail="Percorso non trovato")
    
    # Check if already started
    existing = await db.user_paths.find_one({
        "user_id": user["id"],
        "path_id": path_id
    })
    
    if existing:
        return {"message": "Percorso già iniziato", "path_id": path_id}
    
    # Create path record
    await db.user_paths.insert_one({
        "id": str(uuid.uuid4()),
        "user_id": user["id"],
        "path_id": path_id,
        "started_at": datetime.now(timezone.utc).isoformat(),
        "completed_steps": [],
        "consultations": []
    })
    
    return {"message": "Percorso iniziato", "path_id": path_id}


@api_router.post("/paths/{path_id}/complete-step")
async def complete_path_step(path_id: str, request: Request, step_day: int = 1, consultation_id: str = None):
    """Mark a step as completed"""
    user = await get_current_user(request)
    
    user_path = await db.user_paths.find_one({
        "user_id": user["id"],
        "path_id": path_id
    })
    
    if not user_path:
        raise HTTPException(status_code=404, detail="Percorso non iniziato")
    
    update_data = {
        "$addToSet": {"completed_steps": step_day}
    }
    
    if consultation_id:
        update_data["$push"] = {"consultations": {
            "step_day": step_day,
            "consultation_id": consultation_id,
            "completed_at": datetime.now(timezone.utc).isoformat()
        }}
    
    await db.user_paths.update_one(
        {"user_id": user["id"], "path_id": path_id},
        update_data
    )
    
    # Check if path is now complete and generate synthesis
    path_info = GUIDED_PATHS.get(path_id)
    if path_info:
        total_steps = len(path_info.get("steps", []))
        completed_steps = len(user_path.get("completed_steps", [])) + 1  # Include current step
        
        if completed_steps >= total_steps:
            # Path completed! Generate synthesis
            await generate_path_synthesis(db, user, path_id, path_info)
    
    return {"message": "Passo completato", "step_day": step_day}


async def generate_path_synthesis(db, user, path_id: str, path_info: dict):
    """Generate AI synthesis for completed path"""
    lang = user.get("language", "it")
    
    # Get all consultations for this path
    user_path = await db.user_paths.find_one({
        "user_id": user["id"],
        "path_id": path_id
    })
    
    if not user_path:
        return
    
    consultation_ids = [c["consultation_id"] for c in user_path.get("consultations", [])]
    
    # Fetch all consultations
    consultations = await db.consultations.find({
        "id": {"$in": consultation_ids}
    }).to_list(100)
    
    if not consultations:
        return
    
    # Build context for AI synthesis
    hexagrams_info = []
    for consultation in consultations:
        hex_num = consultation.get("hexagram_number")
        hex_data = get_hexagram_traditional_data(hex_num) if hex_num else {}
        moving_lines = consultation.get("moving_lines", [])
        
        hexagrams_info.append({
            "question": consultation.get("question", ""),
            "hexagram_number": hex_num,
            "hexagram_name": hex_data.get("name", ""),
            "hexagram_meaning": hex_data.get("meaning_it" if lang == "it" else "meaning_en", ""),
            "judgment": hex_data.get("judgment_it" if lang == "it" else "judgment_en", ""),
            "moving_lines": moving_lines,
            "interpretation": consultation.get("interpretation", "")
        })
    
    # Generate AI synthesis
    synthesis_prompt = f"""Sei un saggio maestro dell'I Ching. Un utente ha completato il percorso "{path_info.get('name_it' if lang == 'it' else 'name_en', path_id)}".

Durante il percorso ha consultato l'oracolo con le seguenti domande e ha ricevuto questi esagrammi:

"""
    
    for i, info in enumerate(hexagrams_info, 1):
        synthesis_prompt += f"""
--- Consultazione {i} ---
Domanda: {info['question']}
Esagramma: {info['hexagram_number']} - {info['hexagram_name']}
Significato: {info['hexagram_meaning']}
Giudizio: {info['judgment']}
Linee mutanti: {', '.join(map(str, info['moving_lines'])) if info['moving_lines'] else 'Nessuna'}
"""

    if lang == "it":
        synthesis_prompt += """

Basandoti su TUTTI questi esagrammi e le loro interazioni, crea una SINTESI UNICA e COMPLETA che:

1. **ANALISI COMPLESSIVA**: Identifica il tema centrale che emerge dalla combinazione di tutti gli esagrammi
2. **PUNTI DI FORZA**: Quali qualità e risorse l'utente può sfruttare
3. **AREE DI MIGLIORAMENTO**: Aspetti su cui lavorare per la crescita personale
4. **PIANO D'AZIONE**: Passi concreti e specifici da seguire (minimo 5 punti)
5. **CONSIGLIO FINALE**: Un messaggio di saggezza che integra tutti gli insegnamenti

Scrivi in modo profondo ma accessibile, come un maestro saggio che guida un allievo. Non elencare semplicemente i significati degli esagrammi, ma crea una visione INTEGRATA e PERSONALIZZATA per il percorso di crescita dell'utente.
"""
    else:
        synthesis_prompt += """

Based on ALL these hexagrams and their interactions, create a UNIQUE and COMPLETE SYNTHESIS that:

1. **OVERALL ANALYSIS**: Identify the central theme emerging from the combination of all hexagrams
2. **STRENGTHS**: What qualities and resources the user can leverage
3. **AREAS FOR IMPROVEMENT**: Aspects to work on for personal growth
4. **ACTION PLAN**: Concrete and specific steps to follow (minimum 5 points)
5. **FINAL ADVICE**: A wisdom message that integrates all teachings

Write in a deep but accessible way, like a wise master guiding a student. Don't simply list hexagram meanings, but create an INTEGRATED and PERSONALIZED vision for the user's growth path.
"""
    
    synthesis_text = await _gemini_generate(
        system_instruction="",
        prompt=synthesis_prompt,
        generation_config=GEMINI_DEEP_CONFIG,
    )
    if not synthesis_text or len(synthesis_text) < 100:
        synthesis_text = "Sintesi non disponibile al momento." if lang == "it" else "Synthesis not available at the moment."
    
    # Save completed path with synthesis
    completed_path_doc = {
        "id": str(uuid.uuid4()),
        "user_id": user["id"],
        "path_id": path_id,
        "path_name": path_info.get("name_it" if lang == "it" else "name_en", path_id),
        "path_emoji": path_info.get("emoji", "🎯"),
        "consultations": hexagrams_info,
        "synthesis": synthesis_text,
        "completed_at": datetime.now(timezone.utc).isoformat(),
        "is_read": False  # For notification badge
    }
    
    await db.completed_paths.insert_one(completed_path_doc)
    
    # Update user_path to mark as synthesis_generated
    await db.user_paths.update_one(
        {"user_id": user["id"], "path_id": path_id},
        {"$set": {"synthesis_generated": True, "completed_path_id": completed_path_doc["id"]}}
    )


@api_router.get("/paths/completed")
async def get_completed_paths(request: Request):
    """Get all completed paths with synthesis for current user"""
    user = await get_current_user(request)
    
    completed = await db.completed_paths.find(
        {"user_id": user["id"]},
        {"_id": 0}
    ).sort("completed_at", -1).to_list(100)
    
    return completed


@api_router.get("/paths/completed/{completed_path_id}")
async def get_completed_path_detail(completed_path_id: str, request: Request):
    """Get detail of a specific completed path"""
    user = await get_current_user(request)
    
    completed_path = await db.completed_paths.find_one(
        {"id": completed_path_id, "user_id": user["id"]},
        {"_id": 0}
    )
    
    if not completed_path:
        raise HTTPException(status_code=404, detail="Percorso completato non trovato")
    
    # Mark as read
    if not completed_path.get("is_read"):
        await db.completed_paths.update_one(
            {"id": completed_path_id},
            {"$set": {"is_read": True}}
        )
    
    return completed_path


# ============== PROGRESSION SYSTEM ==============

@api_router.get("/progression")
async def get_user_progression(request: Request):
    """Get user's progression (level, badges, etc.)"""
    user = await get_current_user(request)
    lang = user.get("language", "it")
    
    total = await db.consultations.count_documents({"user_id": user["id"]})
    level_info = get_user_level(total)
    
    # Get user badges
    user_badges = user.get("badges", [])
    badges_detail = []
    for badge in BADGES:
        badge_copy = badge.copy()
        badge_copy["earned"] = badge["id"] in user_badges
        badge_copy["name"] = badge["name_it"] if lang == "it" else badge["name_en"]
        badge_copy["description"] = badge["description_it"] if lang == "it" else badge["description_en"]
        badges_detail.append(badge_copy)
    
    # Format level info
    level_info["current"]["title"] = level_info["current"]["title_it"] if lang == "it" else level_info["current"]["title_en"]
    if level_info["next"]:
        level_info["next"]["title"] = level_info["next"]["title_it"] if lang == "it" else level_info["next"]["title_en"]
    
    return {
        "level": level_info,
        "badges": badges_detail,
        "total_consultations": total
    }


# ============== PERSONALIZED ADVICE SYSTEM (PREMIUM) ==============

class NotificationPreferencesUpdate(BaseModel):
    enabled: Optional[bool] = None
    frequency: Optional[str] = None  # daily, weekly, monthly
    preferred_time: Optional[str] = None  # HH:MM format
    push_enabled: Optional[bool] = None
    in_app_enabled: Optional[bool] = None
    fcm_token: Optional[str] = None


@api_router.get("/advice/daily")
async def get_daily_advice(request: Request):
    """
    Get personalized daily advice based on user's paths and Chinese zodiac calendar.
    PREMIUM ONLY feature.
    """
    user = await get_current_user(request)
    plan = get_user_plan(user)
    
    if plan != "premium":
        raise HTTPException(
            status_code=403, 
            detail="Questa funzionalità è disponibile solo per utenti Premium"
        )
    
    lang = user.get("language", "it")
    advice = await generate_personalized_advice(db, user["id"], "daily", lang)
    
    return advice


@api_router.get("/advice/weekly")
async def get_weekly_advice(request: Request):
    """Get personalized weekly advice. PREMIUM ONLY."""
    user = await get_current_user(request)
    plan = get_user_plan(user)
    
    if plan != "premium":
        raise HTTPException(status_code=403, detail="Funzionalità Premium")
    
    lang = user.get("language", "it")
    advice = await generate_personalized_advice(db, user["id"], "weekly", lang)
    
    return advice


@api_router.get("/advice/monthly")
async def get_monthly_advice(request: Request):
    """Get personalized monthly advice. PREMIUM ONLY."""
    user = await get_current_user(request)
    plan = get_user_plan(user)
    
    if plan != "premium":
        raise HTTPException(status_code=403, detail="Funzionalità Premium")
    
    lang = user.get("language", "it")
    advice = await generate_personalized_advice(db, user["id"], "monthly", lang)
    
    return advice


@api_router.get("/advice/current")
async def get_current_advice(request: Request):
    """
    Get the current advice based on user's notification preference frequency.
    Returns daily/weekly/monthly advice based on settings.
    PREMIUM ONLY.
    """
    user = await get_current_user(request)
    plan = get_user_plan(user)
    
    if plan != "premium":
        # Return limited preview for free users
        day_energy = get_chinese_day_energy()
        year_animal = get_chinese_year_animal()
        lang = user.get("language", "it")
        
        return {
            "is_preview": True,
            "preview_message": "Passa a Premium per ricevere consigli personalizzati basati sui tuoi percorsi!" if lang == "it" else "Upgrade to Premium to receive personalized advice based on your paths!",
            "chinese_calendar": {
                "day_energy": day_energy,
                "year_animal": year_animal,
            }
        }
    
    # Get user's preference
    try:
        prefs = await get_user_notification_preferences(db, user["id"])
        # Strip MongoDB internals (ObjectId is not JSON serializable)
        clean_prefs = {k: v for k, v in prefs.items() if k != "_id"}
        frequency = clean_prefs.get("frequency", "daily")
        lang = user.get("language", "it")

        advice = await generate_personalized_advice(db, user["id"], frequency, lang)
        if not isinstance(advice, dict):
            advice = {"message": str(advice) if advice else ""}
        advice["notification_preferences"] = clean_prefs
        return advice
    except Exception as e:
        logger.error(f"Error generating personalized advice: {e}", exc_info=True)
        # Graceful fallback: return chinese-calendar preview instead of 500
        day_energy = get_chinese_day_energy()
        year_animal = get_chinese_year_animal()
        lang = user.get("language", "it")
        return {
            "is_fallback": True,
            "message": "Consiglio temporaneamente non disponibile. Riprova tra poco." if lang == "it" else "Advice temporarily unavailable. Try again shortly.",
            "chinese_calendar": {
                "day_energy": day_energy,
                "year_animal": year_animal,
            }
        }


@api_router.get("/chinese-calendar")
async def get_chinese_calendar_info(request: Request):
    """Get Chinese calendar information for today (available to all users)"""
    try:
        user = await get_current_user(request)
        lang = user.get("language", "it")
    except:
        lang = "it"
    
    day_energy = get_chinese_day_energy()
    year_animal = get_chinese_year_animal()
    lunar_phase = get_lunar_phase()
    
    return {
        "day_energy": day_energy,
        "year_animal": year_animal,
        "lunar_phase": lunar_phase,
        "date": datetime.now(timezone.utc).isoformat(),
    }


@api_router.get("/notifications/preferences")
async def get_notification_preferences(request: Request):
    """Get user's notification preferences"""
    user = await get_current_user(request)
    prefs = await get_user_notification_preferences(db, user["id"])
    
    # Remove internal fields
    return {
        "enabled": prefs.get("enabled", True),
        "frequency": prefs.get("frequency", "daily"),
        "preferred_time": prefs.get("preferred_time", "08:00"),
        "push_enabled": prefs.get("push_enabled", False),
        "in_app_enabled": prefs.get("in_app_enabled", True),
        "has_fcm_token": bool(prefs.get("fcm_token")),
    }


@api_router.put("/notifications/preferences")
async def update_notification_preferences_endpoint(
    request: Request,
    updates: NotificationPreferencesUpdate
):
    """Update user's notification preferences. PREMIUM ONLY."""
    user = await get_current_user(request)
    plan = get_user_plan(user)
    
    if plan != "premium":
        raise HTTPException(
            status_code=403, 
            detail="Le preferenze di notifica sono disponibili solo per utenti Premium"
        )
    
    updates_dict = updates.dict(exclude_none=True)
    
    # Validate frequency
    if "frequency" in updates_dict and updates_dict["frequency"] not in ["daily", "weekly", "monthly"]:
        raise HTTPException(status_code=400, detail="Frequenza non valida. Usa: daily, weekly, monthly")
    
    # Validate time format
    if "preferred_time" in updates_dict:
        try:
            datetime.strptime(updates_dict["preferred_time"], "%H:%M")
        except ValueError:
            raise HTTPException(status_code=400, detail="Formato ora non valido. Usa: HH:MM")
    
    updated_prefs = await update_user_notification_preferences(db, user["id"], updates_dict)
    
    return {
        "message": "Preferenze aggiornate",
        "preferences": {
            "enabled": updated_prefs.get("enabled", True),
            "frequency": updated_prefs.get("frequency", "daily"),
            "preferred_time": updated_prefs.get("preferred_time", "08:00"),
            "push_enabled": updated_prefs.get("push_enabled", False),
            "in_app_enabled": updated_prefs.get("in_app_enabled", True),
        }
    }


@api_router.post("/notifications/register-push")
async def register_push_token(request: Request, token: str):
    """
    Register FCM token for push notifications.
    This endpoint will be used when Firebase is configured.
    PREMIUM ONLY.
    """
    user = await get_current_user(request)
    plan = get_user_plan(user)
    
    if plan != "premium":
        raise HTTPException(status_code=403, detail="Push notifications are Premium only")
    
    await update_user_notification_preferences(db, user["id"], {
        "fcm_token": token,
        "push_enabled": True
    })
    
    return {"message": "Token registrato con successo", "push_enabled": True}


# ============== USER PROFILE SYSTEM ==============

class UserProfileUpdate(BaseModel):
    birth_date: Optional[str] = None  # YYYY-MM-DD format
    birth_time: Optional[str] = None  # HH:MM format
    birth_place: Optional[str] = None
    gender: Optional[str] = None
    occupation: Optional[str] = None
    iching_experience: Optional[str] = None
    activity_level: Optional[str] = None
    wellness_interests: Optional[List[str]] = None


@api_router.get("/profile")
async def get_user_profile(request: Request):
    """Get user's complete profile including astrological data"""
    user = await get_current_user(request)
    lang = user.get("language", "it")
    
    # Get extended profile from user document
    profile_data = user.get("profile", {})
    
    response = {
        "id": user["id"],
        "name": user.get("name"),
        "email": user.get("email"),
        "language": lang,
        "profile_completed": bool(profile_data.get("birth_date")),
        "profile": profile_data,
        "astrological_profile": None,
    }
    
    # Calculate astrological profile if birth_date is available
    if profile_data.get("birth_date"):
        try:
            from datetime import date
            birth_parts = profile_data["birth_date"].split("-")
            birth_date = date(int(birth_parts[0]), int(birth_parts[1]), int(birth_parts[2]))
            
            astro_profile = get_full_astrological_profile(
                birth_date=birth_date,
                birth_time=profile_data.get("birth_time"),
                birth_place=profile_data.get("birth_place"),
                language=lang
            )
            response["astrological_profile"] = astro_profile
        except Exception as e:
            logger.warning(f"Error calculating astrological profile: {e}")
    
    return response


@api_router.put("/profile")
async def update_user_profile(request: Request, profile_update: UserProfileUpdate):
    """Update user's profile data"""
    user = await get_current_user(request)
    
    # Prepare update data
    update_data = profile_update.dict(exclude_none=True)
    
    # Validate occupation length
    if "occupation" in update_data and len(update_data["occupation"]) > 30:
        raise HTTPException(status_code=400, detail="Occupation must be max 30 characters")
    
    # Validate gender
    valid_genders = ["male", "female", "other", "prefer_not_say"]
    if "gender" in update_data and update_data["gender"] not in valid_genders:
        raise HTTPException(status_code=400, detail="Invalid gender value")
    
    # Validate iching_experience
    valid_experience = ["beginner", "intermediate", "expert"]
    if "iching_experience" in update_data and update_data["iching_experience"] not in valid_experience:
        raise HTTPException(status_code=400, detail="Invalid experience value")
    
    # Validate activity_level
    valid_activity = ["sedentary", "moderate", "active"]
    if "activity_level" in update_data and update_data["activity_level"] not in valid_activity:
        raise HTTPException(status_code=400, detail="Invalid activity level")
    
    # Validate wellness_interests
    valid_interests = ["meditation", "yoga", "taichi", "qigong"]
    if "wellness_interests" in update_data:
        for interest in update_data["wellness_interests"]:
            if interest not in valid_interests:
                raise HTTPException(status_code=400, detail=f"Invalid wellness interest: {interest}")
    
    # Validate date format
    if "birth_date" in update_data:
        try:
            from datetime import datetime
            datetime.strptime(update_data["birth_date"], "%Y-%m-%d")
        except ValueError:
            raise HTTPException(status_code=400, detail="Invalid date format. Use YYYY-MM-DD")
    
    # Validate time format
    if "birth_time" in update_data and update_data["birth_time"]:
        try:
            from datetime import datetime
            datetime.strptime(update_data["birth_time"], "%H:%M")
        except ValueError:
            raise HTTPException(status_code=400, detail="Invalid time format. Use HH:MM")
    
    # Update in database
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    await db.users.update_one(
        {"id": user["id"]},
        {"$set": {"profile": {**user.get("profile", {}), **update_data}}}
    )
    
    # Return updated profile
    return await get_user_profile(request)


@api_router.get("/profile/fields")
async def get_profile_fields(request: Request):
    """Get profile field definitions for form building"""
    try:
        user = await get_current_user(request)
        lang = user.get("language", "it")
    except:
        lang = "it"
    
    # Transform fields for frontend
    fields = []
    for field_name, config in USER_PROFILE_FIELDS.items():
        field = {
            "name": field_name,
            "type": config["type"],
            "required": config["required"],
            "label": config[f"label_{lang}"] if f"label_{lang}" in config else config.get("label_it"),
            "max_length": config.get("max_length"),
        }
        
        if "options" in config:
            field["options"] = [
                {
                    "value": opt["value"],
                    "label": opt[f"label_{lang}"] if f"label_{lang}" in opt else opt.get("label_it")
                }
                for opt in config["options"]
            ]
        
        fields.append(field)
    
    return {"fields": fields}


@api_router.get("/profile/completion-status")
async def get_profile_completion_status(request: Request):
    """Check if user has completed their profile"""
    user = await get_current_user(request)
    profile = user.get("profile", {})
    
    # Check completion
    has_birth_date = bool(profile.get("birth_date"))
    has_basic_info = bool(profile.get("gender") or profile.get("birth_date"))
    
    completion_percentage = 0
    filled_fields = 0
    total_fields = len(USER_PROFILE_FIELDS)
    
    for field_name in USER_PROFILE_FIELDS:
        if profile.get(field_name):
            filled_fields += 1
    
    completion_percentage = int((filled_fields / total_fields) * 100)
    
    return {
        "is_complete": has_birth_date and has_basic_info,
        "completion_percentage": completion_percentage,
        "filled_fields": filled_fields,
        "total_fields": total_fields,
        "missing_essential": not has_birth_date,
        "show_prompt": not has_birth_date,
    }


# ============== NATAL CHART GENERATION ==============

class NatalChartRequest(BaseModel):
    name: Optional[str] = None
    birth_date: str  # YYYY-MM-DD format
    birth_time: str  # HH:MM format
    birth_place: str  # City name or coordinates


@api_router.post("/natal-chart/generate")
async def generate_natal_chart(request: Request, chart_request: NatalChartRequest):
    """
    Generate a complete natal chart with planetary positions, houses, aspects, and SVG diagram.
    Requires birth date, time, and place.
    """
    user = await get_current_user(request)
    lang = user.get("language", "it")
    
    if not KERYKEION_AVAILABLE:
        raise HTTPException(status_code=503, detail="Natal chart service temporarily unavailable")
    
    # Parse birth date
    try:
        birth_parts = chart_request.birth_date.split("-")
        year = int(birth_parts[0])
        month = int(birth_parts[1])
        day = int(birth_parts[2])
    except:
        raise HTTPException(status_code=400, detail="Invalid date format. Use YYYY-MM-DD")
    
    # Parse birth time
    try:
        time_parts = chart_request.birth_time.split(":")
        hour = int(time_parts[0])
        minute = int(time_parts[1])
    except:
        raise HTTPException(status_code=400, detail="Invalid time format. Use HH:MM")
    
    # Geocode birth place
    geo_result = await geocode_location(chart_request.birth_place)
    if not geo_result:
        raise HTTPException(status_code=400, detail="Could not find location. Please try a different city name.")
    
    lat = geo_result["lat"]
    lng = geo_result["lng"]
    city = geo_result.get("display_name", chart_request.birth_place)
    
    # Calculate natal chart
    name = chart_request.name or user.get("name", "User")
    result = calculate_natal_chart(
        name=name,
        year=year,
        month=month,
        day=day,
        hour=hour,
        minute=minute,
        lat=lat,
        lng=lng,
        city=city,
        language=lang
    )
    
    if not result.get("success"):
        raise HTTPException(status_code=500, detail=result.get("error", "Error generating natal chart"))
    
    # Save to user profile
    await db.users.update_one(
        {"id": user["id"]},
        {"$set": {
            "natal_chart": result,
            "natal_chart_generated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    return result


@api_router.get("/natal-chart")
async def get_saved_natal_chart(request: Request):
    """Get user's saved natal chart if available"""
    user = await get_current_user(request)

    natal_chart = user.get("natal_chart")
    if not natal_chart:
        return {"has_chart": False}

    return {
        "has_chart": True,
        "chart": natal_chart,
        "generated_at": user.get("natal_chart_generated_at")
    }


@api_router.get("/natal-chart/svg")
async def get_natal_chart_svg(request: Request):
    """Download natal chart as SVG file"""
    user = await get_current_user(request)
    natal_chart = user.get("natal_chart")
    if not natal_chart or not natal_chart.get("chart_svg"):
        raise HTTPException(status_code=404, detail="Tema natale non trovato")

    name = natal_chart.get("subject", {}).get("name", "tema_natale")
    return Response(
        content=natal_chart["chart_svg"],
        media_type="image/svg+xml",
        headers={"Content-Disposition": f'attachment; filename="tema_natale_{name}.svg"'}
    )


@api_router.get("/natal-chart/pdf")
async def get_natal_chart_pdf(request: Request):
    """Download natal chart as PDF with chart image and details"""
    user = await get_current_user(request)
    natal_chart = user.get("natal_chart")
    if not natal_chart:
        raise HTTPException(status_code=404, detail="Tema natale non trovato")

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak
        from reportlab.lib import colors
        from reportlab.lib.units import cm
        import io
    except ImportError:
        raise HTTPException(status_code=500, detail="reportlab non disponibile")

    subject = natal_chart.get("subject", {})
    name = subject.get("name", "Tema Natale")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=20, textColor=colors.HexColor('#C44D38'), alignment=1)
    h2 = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor('#2C2C2C'))

    story = []
    story.append(Paragraph(f"Tema Natale di {name}", title_style))
    story.append(Spacer(1, 0.5*cm))

    # Subject info
    info = [
        ["Nome", subject.get("name", "")],
        ["Data di nascita", subject.get("birth_date", "")],
        ["Ora di nascita", subject.get("birth_time", "")],
        ["Luogo", subject.get("birth_place", "")],
    ]
    t = Table(info, colWidths=[5*cm, 11*cm])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F9F7F2')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('PADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(t)
    story.append(Spacer(1, 0.5*cm))

    # SVG chart image
    svg_content = natal_chart.get("chart_svg")
    if svg_content:
        try:
            import cairosvg
            svg_resolved = resolve_svg_css_variables(svg_content)
            png_bytes = cairosvg.svg2png(bytestring=svg_resolved.encode("utf-8"), output_width=1000)
            img = Image(io.BytesIO(png_bytes), width=15*cm, height=15*cm)
            story.append(img)
            story.append(PageBreak())
        except Exception as e:
            logger.error(f"Error converting SVG to PNG: {e}")

    # Ascendant
    asc = natal_chart.get("ascendant", {})
    if asc:
        story.append(Paragraph(f"Ascendente: {asc.get('sign','')} {asc.get('degree_formatted','')}", h2))
        story.append(Paragraph(asc.get("interpretation", ""), styles['BodyText']))
        story.append(Spacer(1, 0.3*cm))

    # Planets
    story.append(Paragraph("Pianeti", h2))
    planets = natal_chart.get("planets", [])
    planet_data = [["Pianeta", "Segno", "Grado", "Casa"]]
    for p in planets:
        planet_data.append([
            p.get("name_it", p.get("name", "")),
            p.get("sign", ""),
            p.get("degree_formatted", ""),
            p.get("house", "").replace("_House", "")
        ])
    pt = Table(planet_data, colWidths=[4*cm, 4*cm, 4*cm, 4*cm])
    pt.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#C44D38')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
        ('PADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(pt)
    story.append(Spacer(1, 0.5*cm))

    # Houses
    houses = natal_chart.get("houses", [])
    if houses:
        story.append(Paragraph("Case Astrologiche", h2))
        house_data = [["Casa", "Segno", "Grado"]]
        for h in houses:
            house_data.append([
                str(h.get("number", "")),
                h.get("sign", ""),
                h.get("degree_formatted", "")
            ])
        ht = Table(house_data, colWidths=[3*cm, 6*cm, 6*cm])
        ht.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#7A4F8F')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
            ('PADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(ht)

    doc.build(story)
    buffer.seek(0)

    return Response(
        content=buffer.read(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="tema_natale_{name}.pdf"'}
    )


@api_router.get("/natal-chart/docx")
async def get_natal_chart_docx(request: Request):
    """Download natal chart as DOCX (Word) document"""
    user = await get_current_user(request)
    natal_chart = user.get("natal_chart")
    if not natal_chart:
        raise HTTPException(status_code=404, detail="Tema natale non trovato")

    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx non disponibile")

    subject = natal_chart.get("subject", {})
    name = subject.get("name", "Tema Natale")

    document = Document()

    # Title
    title = document.add_heading(f"Tema Natale di {name}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subject info
    document.add_heading("Dati di Nascita", level=2)
    table = document.add_table(rows=4, cols=2)
    table.style = "Light Grid Accent 1"
    info = [
        ("Nome", subject.get("name", "")),
        ("Data", subject.get("birth_date", "")),
        ("Ora", subject.get("birth_time", "")),
        ("Luogo", subject.get("birth_place", "")),
    ]
    for i, (k, v) in enumerate(info):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v

    # SVG → PNG → image
    svg_content = natal_chart.get("chart_svg")
    if svg_content:
        try:
            import cairosvg
            svg_resolved = resolve_svg_css_variables(svg_content)
            png_bytes = cairosvg.svg2png(bytestring=svg_resolved.encode("utf-8"), output_width=1000)
            document.add_paragraph()
            document.add_picture(io.BytesIO(png_bytes), width=Inches(6))
        except Exception as e:
            logger.error(f"Error converting SVG to PNG: {e}")

    # Ascendant
    asc = natal_chart.get("ascendant", {})
    if asc:
        document.add_heading(f"Ascendente: {asc.get('sign', '')} {asc.get('degree_formatted', '')}", level=2)
        document.add_paragraph(asc.get("interpretation", ""))

    # Planets
    document.add_heading("Pianeti", level=2)
    planets = natal_chart.get("planets", [])
    ptable = document.add_table(rows=len(planets) + 1, cols=4)
    ptable.style = "Light Grid Accent 2"
    headers = ["Pianeta", "Segno", "Grado", "Casa"]
    for i, h in enumerate(headers):
        ptable.cell(0, i).text = h
    for i, p in enumerate(planets, 1):
        ptable.cell(i, 0).text = p.get("name_it", p.get("name", ""))
        ptable.cell(i, 1).text = p.get("sign", "")
        ptable.cell(i, 2).text = p.get("degree_formatted", "")
        ptable.cell(i, 3).text = p.get("house", "").replace("_House", "")

    # Houses
    houses = natal_chart.get("houses", [])
    if houses:
        document.add_heading("Case Astrologiche", level=2)
        htable = document.add_table(rows=len(houses) + 1, cols=3)
        htable.style = "Light Grid Accent 3"
        headers = ["Casa", "Segno", "Grado"]
        for i, h in enumerate(headers):
            htable.cell(0, i).text = h
        for i, h in enumerate(houses, 1):
            htable.cell(i, 0).text = str(h.get("number", ""))
            htable.cell(i, 1).text = h.get("sign", "")
            htable.cell(i, 2).text = h.get("degree_formatted", "")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return Response(
        content=buffer.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="tema_natale_{name}.docx"'}
    )


@api_router.get("/geocode")
async def geocode_city(city: str):
    """Geocode a city name to get coordinates"""
    result = await geocode_location(city)
    if not result:
        raise HTTPException(status_code=404, detail="Location not found")
    return result


# Include the router
app.include_router(api_router)

# CORS: `allow_origins=['*']` combinato con `allow_credentials=True` e'
# vietato dalla spec CORS (alcuni browser rifiutano, ma e' anche un
# CSRF/credential-leak). Default a una whitelist concreta, override via
# env var CORS_ORIGINS per ambienti staging.
_DEFAULT_ORIGINS = (
    "https://chingbenessere.it,"
    "https://www.chingbenessere.it,"
    "http://localhost:3000,"
    "http://localhost:3001,"
    "capacitor://localhost,"          # iOS Capacitor
    "http://localhost,"               # Android Capacitor
    "https://localhost"               # Android Capacitor https scheme
)
_origins = [o.strip() for o in os.environ.get('CORS_ORIGINS', _DEFAULT_ORIGINS).split(',') if o.strip()]
if '*' in _origins:
    # Se davvero serve wildcard, allora `allow_credentials` deve essere False.
    logger.warning("CORS_ORIGINS='*' → disabilito allow_credentials per conformita' alla spec CORS")
    _allow_credentials = False
else:
    _allow_credentials = True

app.add_middleware(
    CORSMiddleware,
    allow_credentials=_allow_credentials,
    allow_origins=_origins,
    allow_methods=["GET", "POST", "PUT", "DELETE", "PATCH", "OPTIONS"],
    allow_headers=["Authorization", "Content-Type", "X-Admin-Secret", "Stripe-Signature"],
)

# Aggancia il rate limiter all'app per gli endpoint marcati con @limiter.limit
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# Gzip everything > 500 bytes. Backend payloads (hexagrams.json, library)
# go from ~30KB to ~6KB. Browser supports gzip out of the box.
app.add_middleware(GZipMiddleware, minimum_size=500, compresslevel=6)


# ──────────────────────────────────────────────────────────────────────
# Security headers — aggiunti su OGNI risposta HTTP.
# ──────────────────────────────────────────────────────────────────────
#   HSTS                        forza HTTPS sui prossimi 6 mesi
#   X-Content-Type-Options      blocca MIME sniffing
#   X-Frame-Options             impedisce clickjacking via iframe
#   Referrer-Policy             non leakkare URL/query a terzi
#   Permissions-Policy          disabilita feature pericolose non usate
#   Cross-Origin-Opener-Policy  isola da popup malevoli (Google Sign-In)
#
# CSP non e' aggiunta dal backend perche' il frontend statico e' servito
# da Vercel — la mettiamo come <meta http-equiv> in index.html.
@app.middleware("http")
async def _security_headers(request: Request, call_next):
    response = await call_next(request)
    # HSTS solo se la richiesta arriva via HTTPS (in dev locale non vogliamo
    # forzare HTTPS che farebbe rompere localhost). Render espone X-Forwarded-Proto.
    proto = request.headers.get("X-Forwarded-Proto") or request.url.scheme
    if proto == "https":
        response.headers.setdefault(
            "Strict-Transport-Security",
            "max-age=15552000; includeSubDomains"  # 180 giorni
        )
    response.headers.setdefault("X-Content-Type-Options", "nosniff")
    response.headers.setdefault("X-Frame-Options", "DENY")
    response.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
    response.headers.setdefault(
        "Permissions-Policy",
        "geolocation=(), microphone=(), camera=(), payment=(), usb=(), "
        "accelerometer=(), gyroscope=()"
    )
    response.headers.setdefault("Cross-Origin-Opener-Policy", "same-origin-allow-popups")
    # Server fingerprinting: nascondiamo la versione FastAPI/Uvicorn
    if response.headers.get("server"):
        response.headers["server"] = "iching"
    return response


# ──────────────────────────────────────────────────────────────────────
# Sentry monitoring — opt-in via SENTRY_DSN env var.
# ──────────────────────────────────────────────────────────────────────
# Se SENTRY_DSN non e' settato, Sentry resta no-op (zero costo runtime).
# Quando l'utente lo configura su Render → ogni eccezione 5xx, errore
# Gemini, errore Stripe, slow query, etc. arriva nella dashboard in
# tempo reale invece di restare sepolto nei log Render (visibili 7 gg).
_sentry_dsn = os.environ.get('SENTRY_DSN')
if _sentry_dsn:
    try:
        import sentry_sdk  # type: ignore
        from sentry_sdk.integrations.fastapi import FastApiIntegration  # type: ignore
        from sentry_sdk.integrations.starlette import StarletteIntegration  # type: ignore
        sentry_sdk.init(
            dsn=_sentry_dsn,
            traces_sample_rate=float(os.environ.get('SENTRY_TRACES_SAMPLE_RATE', '0.1')),
            profiles_sample_rate=float(os.environ.get('SENTRY_PROFILES_SAMPLE_RATE', '0.0')),
            environment=os.environ.get('SENTRY_ENVIRONMENT', 'production'),
            release=os.environ.get('RENDER_GIT_COMMIT', 'unknown')[:7],
            send_default_pii=False,  # GDPR: no PII (email, IP) di default
            integrations=[
                FastApiIntegration(transaction_style='endpoint'),
                StarletteIntegration(transaction_style='endpoint'),
            ],
            ignore_errors=[
                'HTTPException',          # 4xx attesi
                'RateLimitExceeded',      # generati apposta
            ],
        )
        logger.info("✅ Sentry inizializzato (release=%s)", os.environ.get('RENDER_GIT_COMMIT', 'unknown')[:7])
    except ImportError:
        logger.warning("SENTRY_DSN settato ma sentry_sdk non installato — eseguire pip install sentry-sdk[fastapi]")
    except Exception as e:  # noqa: BLE001
        logger.warning("Inizializzazione Sentry fallita: %s", e)


# Endpoints whose payload never changes per-user (truly static catalog data).
# We let the browser cache them for an hour — huge speedup for repeated visits
# to the library and for the consultation flow that reads hex data.
_STATIC_API_PATHS = (
    "/api/hexagrams",
    "/api/library/hexagrams",
    "/api/library/trigrams",
)


@app.middleware("http")
async def add_privacy_and_cache_headers(request: Request, call_next):
    """
    Adds security headers required by the Privacy Policy AND smart cache
    headers per route family:
      - personal endpoints  -> no-store
      - static catalog data -> public, max-age=3600
      - everything else     -> default browser caching
    """
    response = await call_next(request)
    response.headers["X-Robots-Tag"] = "noindex, nofollow, noarchive, nosnippet"
    response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    response.headers["X-Content-Type-Options"] = "nosniff"

    path = request.url.path
    if path.startswith(_STATIC_API_PATHS):
        # 1 hour browser cache + stale-while-revalidate 1 day
        response.headers["Cache-Control"] = "public, max-age=3600, stale-while-revalidate=86400"
    elif path.startswith((
        "/api/auth", "/api/consultations", "/api/profile",
        "/api/natal-chart", "/api/notes", "/api/notifications",
        "/api/fitness", "/api/subscription",
    )):
        response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, private"
    return response


# ────────────────────────────────────────────────────────────────────────
# Performance: MongoDB indexes
# ────────────────────────────────────────────────────────────────────────
@app.on_event("startup")
async def ensure_indexes():
    """
    Create indexes on the collections we read most frequently.
    Idempotent: MongoDB skips creation if the index already exists.
    """
    try:
        await db.users.create_index("email", unique=True)
        await db.users.create_index("id", unique=True)
        await db.consultations.create_index([("user_id", 1), ("created_at", -1)])
        await db.consultations.create_index("share_token")
        await db.notes.create_index([("user_id", 1), ("consultation_id", 1)])
        await db.payment_transactions.create_index("session_id", unique=True)
        await db.payment_transactions.create_index([("user_id", 1), ("created_at", -1)])
        await db.notification_preferences.create_index("user_id", unique=True)
        await db.notification_reads.create_index("user_id", unique=True)
        await db.fitness_profiles.create_index("user_id", unique=True)
        await db.fitness_programs.create_index([("user_id", 1), ("active", 1)])
        await db.fitness_completed.create_index([("user_id", 1), ("completed_at", -1)])
        await db.user_paths.create_index([("user_id", 1), ("path_id", 1)])
        await db.completed_paths.create_index([("user_id", 1), ("is_read", 1)])
        await db.password_resets.create_index([("email", 1), ("created_at", -1)])
        logger.info("✅ MongoDB indexes ensured")
    except Exception as e:
        logger.warning(f"Index creation skipped: {e}")


@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()


# Lightweight ping endpoint used by external uptime monitors (UptimeRobot,
# Cron-job.org, etc.) to keep the Render free-tier service warm and avoid
# the 30-60s cold start when a real user lands on the site.
@app.get("/ping")
@app.get("/healthz")
async def health_check():
    return {"status": "ok", "ts": datetime.now(timezone.utc).isoformat()}
